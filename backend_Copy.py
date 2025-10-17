# backend.py
import uvicorn
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import tempfile, io, os, re, json, base64, hashlib
from typing import List, Tuple, Dict
import fitz  # PyMuPDF
import requests
import pandas as pd
from docx import Document
from io import BytesIO

from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime

from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
import datetime
import os

# Use SQLite database that works anywhere
DATABASE_URL = "sqlite:///./mcq_database.db"

engine = create_engine(
    DATABASE_URL, 
    connect_args={"check_same_thread": False}  # Needed for SQLite
)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class Question(Base):
    __tablename__ = "questions"

    id = Column(Integer, primary_key=True, index=True)
    topic = Column(String(255))
    type = Column(String(20))  # MCQ / Descriptive
    question = Column(Text, nullable=False)
    option_a = Column(Text)
    option_b = Column(Text)
    option_c = Column(Text)
    option_d = Column(Text)
    answer = Column(Text)
    descriptive_answer = Column(Text)
    difficulty = Column(String(10))
    created_at = Column(DateTime, default=datetime.datetime.utcnow)

# Create table if not exists
Base.metadata.create_all(bind=engine)

def save_questions_to_db(results: dict):
    db = SessionLocal()
    try:
        for topic, data in results.items():
            # Save MCQs
            for mcq in data.get("mcqs", []):
                opts = mcq.get("options", [])
                q = Question(
                    topic=topic,
                    type="MCQ",
                    question=mcq.get("question", ""),
                    option_a=opts[0] if len(opts) > 0 else None,
                    option_b=opts[1] if len(opts) > 1 else None,
                    option_c=opts[2] if len(opts) > 2 else None,
                    option_d=opts[3] if len(opts) > 3 else None,
                    answer=mcq.get("answer", ""),
                    difficulty=mcq.get("difficulty", ""),
                    created_at=datetime.datetime.utcnow()
                )
                db.add(q)

            # Save Descriptive
            for dq in data.get("descriptive", []):
                q = Question(
                    topic=topic,
                    type="Descriptive",
                    question=dq.get("question", ""),
                    descriptive_answer=dq.get("answer", ""),
                    difficulty=dq.get("difficulty", ""),
                    created_at=datetime.datetime.utcnow()
                )
                db.add(q)

        db.commit()  # Commit all changes to the database
    except Exception as e:
        db.rollback()  # Rollback in case of an error
        print("❌ DB error:", e)
    finally:
        db.close()


# ---------- CONFIG ----------
OLLAMA_URL = "http://localhost:11434/api/generate"  # change if required
MODEL = "llama3"
HOST = "127.0.0.1"
PORT = 8000


# ---------- FASTAPI ----------
app = FastAPI()



# Path to your design.html
HTML_PATH = r"C:/Users/CITHP/Documents/combined summary video generator/design - Copy.html"

@app.get("/")
async def read_root():
    return FileResponse(HTML_PATH)


app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"], allow_credentials=True)

# Serve static files (put design.html and any assets inside ./static/)
static_dir = os.path.join(os.path.dirname(__file__), "static")
if not os.path.isdir(static_dir):
    os.makedirs(static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Serve design.html at root
@app.get("/", response_class=HTMLResponse)
async def index():
    fpath = os.path.join(static_dir, "design.html")
    if os.path.exists(fpath):
        return HTMLResponse(open(fpath, "r", encoding="utf-8").read())
    return HTMLResponse("<h3>Place design.html inside ./static/ and reload.</h3>")

# ---------- IN-MEMORY STATE & STORE ----------
IN_MEMORY_STORE = {}  # key -> {"data": bytes, "name": str, "mime": str}
STATE = {
    "pdf_uploads": 0,
    "last_pdf_hash": None,
    "last_pdf_pages": 0,
    "mcq_count": 0,
    "desc_count": 0
}

def store_result_bytes(key: str, data: bytes, filename: str, mime: str):
    IN_MEMORY_STORE[key] = {"data": data, "name": filename, "mime": mime}

@app.get("/download/{key}")
async def download_key(key: str):
    item = IN_MEMORY_STORE.get(key)
    if not item:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return StreamingResponse(io.BytesIO(item["data"]), media_type=item["mime"],
                             headers={"Content-Disposition": f"attachment; filename={item['name']}"})

@app.get("/status")
async def status():
    """Return counters for the top dashboard (PDF uploads, pages, counts)."""
    return {
        "pdf_uploads": STATE["pdf_uploads"],
        "last_pdf_pages": STATE["last_pdf_pages"],
        "mcq_count": STATE["mcq_count"],
        "desc_count": STATE["desc_count"]
    }

# ---------- UTIL HELPERS (ported from your Streamlit code) ----------
def clean_text(text: str) -> str:
    if text is None:
        return ""
    return re.sub(r"[\x00-\x1F\x7F]", "", str(text))

def detect_index_range(doc, min_section_hits: int = 3, consecutive_break: int = 2) -> Tuple[int, int]:
    scores = []
    has_contents_flags = []
    for pno in range(doc.page_count):
        try:
            text = doc.load_page(pno).get_text("text") or ""
        except Exception:
            text = ""
        low = text.lower()
        has_contents = bool(re.search(r"\btable of contents\b|\bcontents\b", low))
        count_sections = len(re.findall(r"\b\d{1,2}\.\d+\b", text))
        count_leaders = len(re.findall(r"\.{2,}\s*\d+|\s+\d{1,3}\s*$", text, re.M))
        score = count_sections + 0.6 * count_leaders + (5 if has_contents else 0)
        scores.append(score)
        has_contents_flags.append(has_contents)

    if any(has_contents_flags):
        start_idx = next(i for i, f in enumerate(has_contents_flags) if f)
        end_idx = start_idx
        break_count = 0
        for i in range(start_idx + 1, len(scores)):
            if scores[i] >= 1.0:
                end_idx = i
                break_count = 0
            else:
                break_count += 1
                if break_count >= consecutive_break:
                    break
        return (start_idx + 1, end_idx + 1)

    start_idx = None
    for i, s in enumerate(scores):
        if s >= min_section_hits:
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("Could not auto-detect contents/index pages.")

    end_idx = start_idx
    gap = 0
    for i in range(start_idx + 1, len(scores)):
        if scores[i] >= 1.0:
            end_idx = i
            gap = 0
        else:
            gap += 1
            if gap >= consecutive_break:
                break
    return (start_idx + 1, end_idx + 1)

# ---------- OLLAMA CALLS & PARSERS ----------
def call_ollama(prompt: str, model: str = MODEL, timeout: int = 240) -> str:
    try:
        resp = requests.post(OLLAMA_URL, json={"model": model, "prompt": prompt, "stream": False, "temperature": 0.3, "max_tokens": 800}, timeout=timeout)
        resp.raise_for_status()
        return resp.json().get("response", "") or ""
    except Exception as e:
        return f"OLLAMA_ERROR: {e}"

def generate_mcqs_ollama(topic: str, context: str = "", full_text: str = "", model: str = MODEL, num_qs: int = 5):
    """
    Generate `num_qs` multiple-choice questions with Ollama.
    """
    prompt = f"""
Generate exactly {num_qs} distinct multiple-choice questions for the topic below. For each question include:
- Exactly 4 labeled options A) B) C) D)
- A single-letter correct answer on its own line: Answer: <A/B/C/D>
- (Optional) Difficulty line: Difficulty: <1-5>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Q2. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Continue this pattern for exactly {num_qs} questions.

Topic: {topic}
Context: {context[:1500]}
"""
    out = call_ollama(prompt, model=model)
    if out.startswith("OLLAMA_ERROR"):
        # If Ollama fails, generate fallback questions to match the requested count
        return generate_fallback_mcqs(topic, num_qs, context)
    
    mcqs = []
    # split by Qn blocks
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = [ln.rstrip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        
        q_line_idx = 0
        for idx, ln in enumerate(lines):
            if re.match(r'^\s*Q\d+\.', ln, re.I):
                q_line_idx = idx
                break
            if not re.match(r'^[A-D][\)\.\-:]', ln, re.I) and not re.search(r'(here are|multiple[-\s]?choice|based on the topic)', ln, re.I):
                q_line_idx = idx
                break
        
        q_line = clean_text(lines[q_line_idx])
        opts = []
        opt_end_idx = q_line_idx
        
        for j in range(q_line_idx + 1, len(lines)):
            m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
            if m:
                label = m.group(1).upper()
                text = m.group(2).strip()
                opts.append(f"{label}. {text}")
                opt_end_idx = j
            else:
                # handle continuation lines for previous option (concatenate)
                if opts and lines[j].strip():
                    opts[-1] = opts[-1] + " " + lines[j].strip()
                else:
                    break
        
        answer = ""
        difficulty = ""
        look_start = opt_end_idx + 1
        look_end = min(len(lines), opt_end_idx + 8)
        
        for k in range(look_start, look_end):
            ln = lines[k]
            m_ans = re.search(r'(?i)\b(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', ln)
            if m_ans:
                answer = m_ans.group(1).upper()
                continue
            m_diff = re.search(r'(?i)\b(?:difficulty|level)[:\s\-]*\(?\s*([1-5])\s*\)?', ln)
            if m_diff:
                difficulty = m_diff.group(1)
                continue
            m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', ln, re.I)
            if m_single and not answer:
                answer = m_single.group(1).upper()
        
        if not answer:
            m_any = re.search(r'(?i)\banswer[:\s\-]*\(?\s*([A-D])\s*\)?', block)
            if m_any:
                answer = m_any.group(1).upper()
        
        if q_line and len(opts) >= 2:
            mcqs.append({"question": q_line, "options": opts, "answer": answer, "difficulty": difficulty})
    
    # Ensure we have exactly the requested number of questions
    if len(mcqs) < num_qs:
        # Generate additional fallback questions to reach the target count
        additional_needed = num_qs - len(mcqs)
        fallback_mcqs = generate_fallback_mcqs(topic, additional_needed, context)
        mcqs.extend(fallback_mcqs)
    elif len(mcqs) > num_qs:
        # Trim excess questions
        mcqs = mcqs[:num_qs]
    
    return mcqs

def generate_fallback_mcqs(topic: str, num_qs: int, context: str = ""):
    """Generate simple fallback MCQs when Ollama fails or returns insufficient questions."""
    mcqs = []
    for i in range(num_qs):
        mcq = {
            "question": f"What is the main concept of '{topic}'?",
            "options": [
                "A. Fundamental principle discussed in the text",
                "B. Basic terminology introduction", 
                "C. Key application mentioned",
                "D. Core methodology described"
            ],
            "answer": "A",
            "difficulty": "3"
        }
        mcqs.append(mcq)
    return mcqs

def generate_descriptive_with_answers(topic: str, context: str = "", model: str = MODEL, num_qs: int = 3):
    prompt = f"""
Generate exactly {num_qs} descriptive / short-answer / essay-style questions for the topic below.
For each question, also provide:
- Correct answer
- Difficulty level (1-5)

Return exactly in this format:

Q1. <question text>
Answer: <answer text>
Difficulty: <1-5>

Q2. <question text>
Answer: <answer text>
Difficulty: <1-5>

Continue this pattern for exactly {num_qs} questions.

Do not add extra commentary.

Topic: {topic}
Context: {context[:1500]}
"""
    out = call_ollama(prompt, model=model)
    if out.startswith("OLLAMA_ERROR"):
        return generate_fallback_descriptive(topic, num_qs, context)
    
    blocks = re.split(r'\n(?=Q\d+\.)', out)
    results = []
    
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        question = ""
        answer = ""
        difficulty = ""
        
        for ln in lines:
            ln = ln.strip()
            if ln.lower().startswith("q"):
                question = re.sub(r'^q\d+\.\s*', '', ln, flags=re.I).strip()
            elif ln.lower().startswith("answer:"):
                answer = ln.split(":", 1)[1].strip()
            elif ln.lower().startswith("difficulty:"):
                difficulty = ln.split(":", 1)[1].strip()
        
        if question:
            results.append({"question": question, "answer": answer, "difficulty": difficulty})
    
    # Ensure we have exactly the requested number of questions
    if len(results) < num_qs:
        additional_needed = num_qs - len(results)
        fallback_desc = generate_fallback_descriptive(topic, additional_needed, context)
        results.extend(fallback_desc)
    elif len(results) > num_qs:
        results = results[:num_qs]
    
    return results

def generate_fallback_descriptive(topic: str, num_qs: int, context: str = ""):
    """Generate simple fallback descriptive questions."""
    results = []
    for i in range(num_qs):
        result = {
            "question": f"Explain the key aspects of '{topic}' as discussed in the text.",
            "answer": f"The text discusses various aspects of {topic} including fundamental concepts, applications, and methodologies.",
            "difficulty": "3"
        }
        results.append(result)
    return results

def build_docx_bytes(questions_data: dict) -> bytes:
    doc = Document()
    doc.add_heading("Generated Questions", level=1)
    for topic_title, blocks in questions_data.items():
        doc.add_heading(topic_title, level=2)
        mcqs = blocks.get("mcqs", []) or []
        if mcqs:
            doc.add_paragraph("Multiple Choice Questions:")
            for idx, mcq in enumerate(mcqs, start=1):
                doc.add_paragraph(f"{idx}. {mcq.get('question','')}")
                for opt in mcq.get("options", []):
                    doc.add_paragraph(f"    {opt}")
                ans = mcq.get("answer", "")
                diff = mcq.get("difficulty", "N/A")
                if ans:
                    doc.add_paragraph(f"    Answer: {ans}    Difficulty: {diff}")
                else:
                    doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
        descrs = blocks.get("descriptive", []) or []
        if descrs:
            doc.add_paragraph("Descriptive / Short-answer Questions:")
            for idx, dq in enumerate(descrs, start=1):
                if isinstance(dq, dict):
                    q = dq.get("question", "")
                    a = dq.get("answer", "")
                    diff = dq.get("difficulty", "N/A")
                else:
                    q = str(dq)
                    a, diff = "", "N/A"
                doc.add_paragraph(f"{idx}. {q}")
                if a:
                    doc.add_paragraph(f"    Answer: {a}")
                doc.add_paragraph(f"    Difficulty: {diff}")
                doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_dfs_from_questions(questions_data: dict):
    rows = []
    for topic_title, topic_data in questions_data.items():
        for mcq in topic_data.get("mcqs", []):
            opts = mcq.get("options") or []
            rows.append({
                "Topic": topic_title,
                "Type": "MCQ",
                "Question": mcq.get("question", ""),
                "Option A": opts[0] if len(opts) > 0 else "",
                "Option B": opts[1] if len(opts) > 1 else "",
                "Option C": opts[2] if len(opts) > 2 else "",
                "Option D": opts[3] if len(opts) > 3 else "",
                "Answer": mcq.get("answer", ""),
                "Difficulty": mcq.get("difficulty", "N/A"),
                "Descriptive Answer": ""
            })
        for dq in topic_data.get("descriptive", []):
            rows.append({
                "Topic": topic_title,
                "Type": "Descriptive",
                "Question": dq.get("question", ""),
                "Option A": "", "Option B": "", "Option C": "", "Option D": "",
                "Answer": "",
                "Difficulty": dq.get("difficulty", "N/A"),
                "Descriptive Answer": dq.get("answer", "")
            })
    return pd.DataFrame(rows)

# ---------- ENDPOINTS: PDF / TOC / GENERATION ----------
@app.post("/extract_toc")
async def extract_toc(file: UploadFile = File(...)):
    pdf_bytes = await file.read()
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # update page count state (not counting as upload until generation)
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        # Try detect TOC pages and parse numeric headings
        try:
            start, end = detect_index_range(doc)
        except Exception:
            start, end = 1, min(6, doc.page_count)
        text = "\n".join([doc.load_page(p-1).get_text("text") or "" for p in range(start, end+1)])
        raw_matches = re.findall(r"(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\b", text)
        matches = []
        if raw_matches:
            for num, title, pno in raw_matches:
                title_clean = re.sub(r"\.{2,}|\.{3,}", ".", title).strip(' .\t')
                title_clean = clean_text(title_clean)
                page_no = int(pno) if pno.isdigit() else None
                matches.append({"subnum": num.strip(), "title": title_clean, "page": page_no})
        else:
            # fallback: search simple lines
            for ln in text.splitlines():
                m = re.match(r'^\s*(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\s*$', ln)
                if m:
                    matches.append({"subnum": m.group(1), "title": clean_text(m.group(2).strip()), "page": int(m.group(3))})
        # Build chapters map
        chapters = {}
        for m in matches:
            chap = int(m["subnum"].split(".")[0]) if m["subnum"].split(".")[0].isdigit() else 0
            chapters.setdefault(chap, []).append(m)
        return {"status": "success", "matches": matches, "chapters_count": len(chapters), "pages": STATE["last_pdf_pages"]}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/generate_pdf_mcqs")
async def generate_pdf_mcqs(
    file: UploadFile = File(...),
    chapters: str = Form("[]"),
    question_type: str = Form("both"),   # "mcq", "descriptive", or "both"
    mcq_source: str = Form("llama_open"), # currently unused by backend, kept for future use
    num_mcqs: int = Form(5),              # Number of MCQs per topic
    num_desc: int = Form(3)               # Number of descriptive questions per topic
):
    pdf_bytes = await file.read()
    selected_chapters = json.loads(chapters)
    qtype = (question_type or "both").lower()
    
    try:
        
        md5 = hashlib.md5(pdf_bytes).hexdigest()
        if STATE.get("last_pdf_hash") != md5:
            STATE["pdf_uploads"] += 1
            STATE["last_pdf_hash"] = md5

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        STATE["last_pdf_pages"] = getattr(doc, "page_count", 0)
        full_text = "\n".join([doc.load_page(p).get_text("text") or "" for p in range(doc.page_count)])

        try:
            start, end = detect_index_range(doc)
            index_text = "\n".join([doc.load_page(p-1).get_text("text") or "" for p in range(start, end+1)])
        except Exception:
            index_text = full_text[:4000]

        raw_matches = re.findall(r"(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\b", index_text)
        topics = []
        if raw_matches:
            for num, title, pno in raw_matches:
                title_clean = clean_text(re.sub(r"\.{2,}|\.{3,}", ".", title).strip(' .\t'))
                page_no = int(pno) if pno.isdigit() else None
                topics.append({"subnum": num, "title": title_clean, "page": page_no})
        else:
            for ln in index_text.splitlines():
                m = re.match(r'^\s*(\d{1,2}\.\d+)\s+(.+)$', ln)
                if m:
                    topics.append({"subnum": m.group(1), "title": clean_text(m.group(2).strip()), "page": None})

        # Filter by selected chapters if provided
        if selected_chapters:
            filtered = []
            for t in topics:
                chap_no = int(t["subnum"].split(".")[0]) if t["subnum"].split(".")[0].isdigit() else 0
                if chap_no in selected_chapters:
                    filtered.append(t)
            topics = filtered

        # Decide which types to produce
        produce_mcq = (qtype in ("mcq", "both"))
        produce_desc = (qtype in ("descriptive", "both"))

        # Generate questions for each topic (only requested types)
        results = {}
        total_mcqs_generated = 0
        total_desc_generated = 0
        
        for t in topics:
            title = t["title"]
            if t.get("page"):
                pg = t["page"]
                startp = max(0, pg-2)
                endp = min(doc.page_count, pg+1)
                context = "\n".join([doc.load_page(p).get_text("text") or "" for p in range(startp, endp)])
            else:
                context = index_text[:2000]

            entry = {}
            if produce_mcq:
                # Use the user-specified number of MCQs
                entry["mcqs"] = generate_mcqs_ollama(title, context=context, full_text=full_text, num_qs=num_mcqs)
                total_mcqs_generated += len(entry["mcqs"])
            else:
                entry["mcqs"] = []

            if produce_desc:
                # Use the user-specified number of descriptive questions
                entry["descriptive"] = generate_descriptive_with_answers(title, context=context, num_qs=num_desc)
                total_desc_generated += len(entry["descriptive"])
            else:
                entry["descriptive"] = []

            results[title] = entry

        # Save the generated questions to the database
        save_questions_to_db(results)

        # Build files and store them
        df_all = build_dfs_from_questions(results)

        # CSV
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "questions.csv", "text/csv")

        # Excel
        excel_buf = BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            df_all.to_excel(writer, sheet_name="Questions", index=False)
        excel_buf.seek(0)
        excel_bytes = excel_buf.getvalue()
        excel_key = hashlib.md5(excel_bytes).hexdigest()
        store_result_bytes(excel_key, excel_bytes, "questions.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX
        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(docx_key, docx_bytes, "questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Update global state with exact counts
        STATE["mcq_count"] = STATE.get("mcq_count", 0) + total_mcqs_generated
        STATE["desc_count"] = STATE.get("desc_count", 0) + total_desc_generated

        return {
            "status": "success",
            "results_count_topics": len(results),
            "mcqCount": total_mcqs_generated,  # Exact count of MCQs generated
            "descCount": total_desc_generated,  # Exact count of descriptive questions generated
            "download_keys": {"csv": csv_key, "excel": excel_key, "docx": docx_key},
            "pages": STATE["last_pdf_pages"],
            "global_state": {
                "pdf_uploads": STATE["pdf_uploads"],
                "last_pdf_pages": STATE["last_pdf_pages"],
                "mcq_count": STATE["mcq_count"],
                "desc_count": STATE["desc_count"]
            },
            "results": results,  # for immediate front-end rendering
            "requested_mcqs_per_topic": num_mcqs,  # For debugging
            "requested_desc_per_topic": num_desc    # For debugging
        }

       

    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.get("/questions")
def get_questions(search: str = None, qtype: str = None):
    db = SessionLocal()
    try:
        query = db.query(Question)
        
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                Question.question.ilike(search_term) |
                Question.topic.ilike(search_term) |
                Question.option_a.ilike(search_term) |
                Question.option_b.ilike(search_term) |
                Question.option_c.ilike(search_term) |
                Question.option_d.ilike(search_term) |
                Question.answer.ilike(search_term) |
                Question.descriptive_answer.ilike(search_term)
            )
        
        if qtype:
            query = query.filter(Question.type == qtype)
            
        questions = query.order_by(Question.created_at.desc()).all()
        
        # Convert to dict for JSON serialization
        result = []
        for q in questions:
            result.append({
                "id": q.id,
                "topic": q.topic,
                "type": q.type,
                "question": q.question,
                "option_a": q.option_a,
                "option_b": q.option_b,
                "option_c": q.option_c,
                "option_d": q.option_d,
                "answer": q.answer,
                "descriptive_answer": q.descriptive_answer,
                "difficulty": q.difficulty,
                "created_at": q.created_at.isoformat() if q.created_at else None
            })
            
        return result
        
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    finally:
        db.close()


@app.post("/save_questions_to_db")
async def save_questions_to_db_endpoint(data: dict):
    try:
        save_questions_to_db(data)  # Calling the existing function to save questions to DB
        return JSONResponse(content={"status": "success"})
    except Exception as e:
        return JSONResponse(content={"status": "error", "error": str(e)}, status_code=500)


from random import sample
# @app.post("/generate_question_paper")
# async def generate_question_paper(levels: Dict[int, int]):
#     """
#     Generate a question paper with random questions based on the selected levels and number of questions per level.
#     """
#     db = SessionLocal()
#     try:
#         questions_by_level = {1: [], 2: [], 3: [], 4: [], 5: []}
        
#         # Fetch questions from the database and group by difficulty level and type
#         all_questions = db.query(Question).filter(Question.type.in_(['MCQ', 'Descriptive'])).all()  # Fetch both MCQ and Descriptive questions
        
#         for q in all_questions:
#             if q.difficulty and q.difficulty.isdigit():
#                 level = int(q.difficulty)
#                 if 1 <= level <= 5:
#                     questions_by_level[level].append(q)
        
#         # Create a paper by selecting random questions from each level
#         question_paper = []
#         total_selected = 0
        
#         for level, count in levels.items():
#             if count > 0 and level in questions_by_level:
#                 available_questions = questions_by_level[level]
#                 if available_questions:
#                     # Select min(count, available) questions randomly
#                     num_to_select = min(count, len(available_questions))
#                     selected_questions = sample(available_questions, num_to_select)
#                     question_paper.extend(selected_questions)
#                     total_selected += num_to_select
        
#         # Return the selected question paper data
#         paper_data = []
#         for q in question_paper:
#             paper_data.append({
#                 "id": q.id,
#                 "topic": q.topic,
#                 "type": q.type,
#                 "question": q.question,
#                 "option_a": q.option_a,
#                 "option_b": q.option_b,
#                 "option_c": q.option_c,
#                 "option_d": q.option_d,
#                 # "descriptive_answer": q.descriptive_answer,
#                 "difficulty": q.difficulty
#             })
        
#         # If the question type is MCQ or Descriptive, do not include the answer
#         if q.type == 'MCQ':
#             # Exclude answer
#             del paper_data['Answer:']

        
#         return {
#             "status": "success", 
#             "questions": paper_data,
#             "total_selected": total_selected,
#             "message": f"Generated paper with {total_selected} questions"
#         }
    
#     except Exception as e:
#         return {"status": "error", "error": str(e)}
#     finally:
#         db.close()





import re
from random import sample

@app.post("/generate_question_paper")
async def generate_question_paper(levels: Dict[int, int]):
    """
    Generate a question paper with random questions based on the selected levels and number of questions per level.
    """
    db = SessionLocal()
    try:
        questions_by_level = {1: [], 2: [], 3: [], 4: [], 5: []}
        
        # Fetch questions from the database and group by difficulty level and type
        all_questions = db.query(Question).filter(Question.type.in_(['MCQ', 'Descriptive'])).all()
        
        for q in all_questions:
            if q.difficulty and q.difficulty.isdigit():
                level = int(q.difficulty)
                if 1 <= level <= 5:
                    questions_by_level[level].append(q)
        
        # Create a paper by selecting random questions from each level
        question_paper = []
        total_selected = 0
        
        for level, count in levels.items():
            if count > 0 and level in questions_by_level:
                available_questions = questions_by_level[level]
                if available_questions:
                    num_to_select = min(count, len(available_questions))
                    selected_questions = sample(available_questions, num_to_select)
                    question_paper.extend(selected_questions)
                    total_selected += num_to_select
        
        # Return the selected question paper data
        paper_data = []
        for q in question_paper:
            # Clean the options to remove answer and difficulty info
            def clean_option(option_text):
                if not option_text:
                    return option_text
                
                # Remove "Answer: X Difficulty: Y" patterns from options
                option_text = re.sub(r'\s*Answer:\s*[A-D]\s*Difficulty:\s*\d\s*$', '', option_text, flags=re.IGNORECASE)
                option_text = re.sub(r'\s*Difficulty:\s*\d\s*Answer:\s*[A-D]\s*$', '', option_text, flags=re.IGNORECASE)
                
                # Remove standalone patterns
                option_text = re.sub(r'\s*Answer:\s*[A-D]\s*$', '', option_text, flags=re.IGNORECASE)
                option_text = re.sub(r'\s*Difficulty:\s*\d\s*$', '', option_text, flags=re.IGNORECASE)
                
                # Final cleanup
                option_text = re.sub(r'[\.\s]*$', '', option_text).strip()
                return option_text

            # Add sanitized question to the result
            question_dict = {
                "id": q.id,
                "topic": q.topic,
                "type": q.type,
                "question": q.question.strip(),  # Question text is already clean
                "option_a": clean_option(q.option_a),
                "option_b": clean_option(q.option_b),
                "option_c": clean_option(q.option_c),
                "option_d": clean_option(q.option_d),
                "difficulty": q.difficulty
            }

            paper_data.append(question_dict)
        
        return {
            "status": "success", 
            "questions": paper_data,
            "total_selected": total_selected,
            "message": f"Generated paper with {total_selected} questions"
        }
    
    except Exception as e:
        return {"status": "error", "error": str(e)}
    finally:
        db.close()






# ---------- RUN ----------
if __name__ == "__main__":
    print(f"Starting backend at http://{HOST}:{PORT}  (static files dir: {static_dir})")
    uvicorn.run("backend:app", host=HOST, port=PORT, reload=True)









# --- VIDEO / TRANSCRIPTION / SUMMARY / MCQ FLOW (paste into backend.py) ---
# required imports at top of file (if not already present)
# --- VIDEO / TRANSCRIPTION / SUMMARY / MCQ FLOW (paste into backend.py) ---
# required imports at top of file (if not already present)
import nltk
from nltk.tokenize import sent_tokenize
try:
    nltk.download('punkt', quiet=True)
except Exception:
    pass

# optional libs flags
try:
    import whisper
    _HAS_WHISPER = True
except Exception:
    _HAS_WHISPER = False

try:
    from moviepy.editor import VideoFileClip
    _HAS_MOVIEPY = True
except Exception:
    _HAS_MOVIEPY = False

# summarizer config (BART chunking)
CHUNK_WORDS = 800
SUMMARIZER_MODEL = "facebook/bart-large-cnn"
SUMMARY_MIN_LENGTH = 30

# Local summarizer via transformers (optional, heavy)
def split_transcript_into_chunks_by_words(transcript: str, chunk_words: int = CHUNK_WORDS):
    sentences = sent_tokenize(transcript)
    chunks, current, current_words = [], [], 0
    for s in sentences:
        wcount = len(s.split())
        if current_words + wcount > chunk_words and current:
            chunks.append(" ".join(current))
            current, current_words = [s], wcount
        else:
            current.append(s)
            current_words += wcount
    if current:
        chunks.append(" ".join(current))
    return chunks

def summarizer_pipeline(model_name=SUMMARIZER_MODEL):
    try:
        from transformers import pipeline
        return pipeline("summarization", model=model_name, device=-1)  # CPU
    except Exception:
        return None

def summarize_chunks(chunks, summarizer):
    summaries = []
    for c in chunks:
        if summarizer:
            try:
                out = summarizer(c, max_length=400, min_length=100, do_sample=False)
                summary_text = out[0]['summary_text'].strip()
            except Exception:
                summary_text = " ".join(c.split()[:SUMMARY_MIN_LENGTH])
        else:
            # fallback: truncate
            summary_text = " ".join(c.split()[:SUMMARY_MIN_LENGTH])
        summaries.append(summary_text)
    return summaries

def combine_and_summarize_summaries(summaries):
    if not summaries:
        return ""
    return "\n\n".join(summaries)

def summarize_transcript_with_bart(transcript: str):
    """
    Try to summarize transcript using local BART in chunks; if local summarizer not available,
    return empty chunks and caller should fallback to Ollama summarizer with summarize_text().
    """
    if not transcript or not transcript.strip():
        return {"overall": "", "chunks": []}
    chunks = split_transcript_into_chunks_by_words(transcript, CHUNK_WORDS)
    summarizer = summarizer_pipeline(SUMMARIZER_MODEL)
    if summarizer is None:
        # signal to caller that local summarizer isn't available
        return {"overall": "", "chunks": []}
    chunk_summaries = summarize_chunks(chunks, summarizer)
    overall_summary = combine_and_summarize_summaries(chunk_summaries)
    return {"overall": overall_summary, "chunks": chunk_summaries}

# Robust MCQ parser (accepts many model output formats)
def parse_mcqs_freeform(output: str) -> List[Dict]:
    mcqs = []
    if not output:
        return mcqs
    raw_lines = [ln.rstrip() for ln in output.splitlines() if ln.strip()]
    # drop very generic intro / header-only lines
    lines = []
    for ln in raw_lines:
        if re.search(r"(here are|multiple[-\s]?choice questions|based on the summary|based on the topic|following questions|the following)", ln, re.I):
            continue
        if re.match(r'^\s*(?:question|q)\s*\d+\b[:.\s-]*$', ln, re.I):
            continue
        lines.append(ln.strip())

    i = 0
    while i < len(lines):
        ln = lines[i]
        # skip stray option lines until we find a question
        if re.match(r'^[A-D][\)\.\-:]\s+', ln, re.I):
            i += 1
            continue
        question_text = re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', ln, flags=re.I).strip()
        if len(question_text) < 3:
            i += 1
            continue
        # collect options
        opts = []
        opt_map = {}
        j = i + 1
        while j < len(lines) and len(opts) < 4:
            if re.match(r'^[A-D][\)\.\-:]\s+', lines[j], re.I):
                m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
                if m:
                    label = m.group(1).upper()
                    text = m.group(2).strip()
                    formatted = f"{label}. {text}"
                    opts.append(formatted)
                    opt_map[label] = formatted
                else:
                    opts.append(lines[j].strip())
                j += 1
            else:
                break
        # look ahead for Answer:
        answer = ""
        look_end = min(len(lines), j + 6)
        for k in range(j, look_end):
            candidate = lines[k].strip()
            m_ans = re.match(r'(?i)^\s*(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', candidate)
            if m_ans:
                answer = m_ans.group(1).upper()
                break
            m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', candidate, re.I)
            if m_single:
                answer = m_single.group(1).upper()
                break
        if answer and answer not in opt_map:
            answer = ""  # validate
        if question_text and len(opts) >= 2:
            mcqs.append({"question": question_text, "options": opts, "answer": answer})
        i = j if j > i else i + 1
    return mcqs
# whisper-based transcription (uses whisper library, raises if not installed)
def split_audio(audio_path: str, chunk_length_sec: int = 300):
    try:
        from pydub import AudioSegment
    except Exception:
        return [audio_path]
    import wave, contextlib
    with contextlib.closing(wave.open(audio_path, 'rb')) as wf:
        rate = wf.getframerate()
        n_frames = wf.getnframes()
        total_sec = n_frames / float(rate)
    if total_sec <= chunk_length_sec:
        return [audio_path]
    audio = AudioSegment.from_wav(audio_path)
    chunk_files = []
    for start_ms in range(0, len(audio), chunk_length_sec * 1000):
        chunk = audio[start_ms:start_ms + chunk_length_sec * 1000]
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        chunk.export(tmp.name, format="wav")
        chunk_files.append(tmp.name)
    return chunk_files

def transcribe_video_bytes(video_bytes: bytes, whisper_model_name: str = "small") -> str:
    if not _HAS_WHISPER or not _HAS_MOVIEPY:
        raise RuntimeError("Whisper or moviepy not available on server.")
    # write video to temp file
    vf = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
    vf.write(video_bytes); vf.flush(); vf.close()
    audio_path = None
    try:
        clip = VideoFileClip(vf.name)
        af = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        audio_path = af.name
        clip.audio.write_audiofile(audio_path, logger=None)
        clip.close()
        chunk_files = split_audio(audio_path)
        model = whisper.load_model(whisper_model_name)
        full_text = ""
        for c in chunk_files:
            res = model.transcribe(c)
            text = res.get("text", "").strip()
            if text:
                full_text += text + " "
            try:
                if c != audio_path and os.path.exists(c):
                    os.remove(c)
            except Exception:
                pass
        return full_text.strip()
    finally:
        try:
            if os.path.exists(vf.name): os.remove(vf.name)
        except Exception:
            pass
        try:
            if audio_path and os.path.exists(audio_path): os.remove(audio_path)
        except Exception:
            pass

# generate MCQs from summary (reuse existing function if present)
def generate_mcqs_from_summary_local(summary: str, num_qs: int = 10, model: str = MODEL):
    # Reuse the same approach as your Streamlit function generate_mcqs_from_summary
    prompt = f"""
Generate {num_qs} distinct multiple-choice questions that cover the following summary.
For each question include:
- Exactly 4 labeled options A) B) C) D) 
- A single-letter answer line like: Answer: <A/B/C/D>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <A/B/C/D>

Summary:
{summary}
"""
    out = call_ollama(prompt, model=model, timeout=600)
    if out.startswith("OLLAMA_ERROR"):
        return [{"question": out, "options": [], "answer": ""}]
    return parse_mcqs_freeform(out)

# Endpoint: transcribe -> summarize (video)
@app.post("/transcribe_video")
async def transcribe_video(file: UploadFile = File(...), whisper_model: str = Form("small")):
    """
    Accepts a video file and returns transcript + summary.
    If local BART summarizer (transformers) is available it will be used; otherwise Ollama summarization used.
    """
    video_bytes = await file.read()
    try:
        # Transcribe (Whisper)
        if not _HAS_WHISPER or not _HAS_MOVIEPY:
            return {"status": "error", "error": "Transcription requires whisper and moviepy installed on server."}
        # update unique-video counter
        try:
            md5 = hashlib.md5(video_bytes).hexdigest()
            if STATE.get("last_video_hash") != md5:
                STATE["video_uploads"] = STATE.get("video_uploads", 0) + 1
                STATE["last_video_hash"] = md5
        except Exception:
            pass
        transcript = transcribe_video_bytes(video_bytes, whisper_model_name=whisper_model)
        # Try local BART summarizer first
        summ = summarize_transcript_with_bart(transcript)
        if not summ["overall"]:
            # fallback: use Ollama summarizer (summarize_text uses Ollama)
            overall = summarize_text(transcript, model=MODEL, max_words=200)
            return {"status": "success", "transcript": transcript, "summary": overall, "chunks": summ["chunks"]}
        return {"status": "success", "transcript": transcript, "summary": summ["overall"], "chunks": summ["chunks"],"global_state": {
                "video_uploads": STATE.get("video_uploads", 0),}}
    except Exception as e:
        return {"status": "error", "error": str(e)}

# Endpoint: generate MCQs (from summary or from video file)
@app.post("/generate_video_mcqs")
async def generate_video_mcqs(
    file: UploadFile = File(None),
    summary: str = Form(""),
    question_type: str = Form("both"),   # "mcq", "descriptive", "both"
    num_qs: int = Form(10),
    whisper_model: str = Form("small")
):
    """
    Generate MCQs (and optionally descriptive questions) from a provided summary string,
    or from an uploaded video file (which will be transcribed & summarized).
    Returns per-request counts and download keys.
    """
    qtype = (question_type or "both").lower()
    summary_text = summary or ""
    try:
        # If file provided and summary empty, transcribe & summarize first
        if file is not None and not summary_text:
            if not _HAS_WHISPER or not _HAS_MOVIEPY:
                return {"status": "error", "error": "Transcription requires whisper and moviepy installed on server."}
            video_bytes = await file.read()
            transcript = transcribe_video_bytes(video_bytes, whisper_model_name=whisper_model)
            # try local BART
            summ = summarize_transcript_with_bart(transcript)
            if summ["overall"]:
                summary_text = summ["overall"]
                chunk_summaries = summ["chunks"]
            else:
                # fallback to Ollama
                summary_text = summarize_text(transcript, model=MODEL, max_words=200)
                chunk_summaries = summ["chunks"]
        elif summary_text:
            chunk_summaries = []
        else:
            return {"status": "error", "error": "No summary or file provided."}

        produce_mcq = (qtype in ("mcq", "both"))
        produce_desc = (qtype in ("descriptive", "both"))

        results = {}
        # We'll treat this as single topic "Video Summary"
        if produce_mcq:
            mcqs = generate_mcqs_from_summary_local(summary_text, num_qs=num_qs, model=MODEL)
        else:
            mcqs = []
        if produce_desc:
            descrs = generate_descriptive_with_answers("Video summary", context=summary_text, model=MODEL, num_qs=3)
        else:
            descrs = []

        results["Video summary"] = {"mcqs": mcqs, "descriptive": descrs}

        # Build files only containing the selected types
        df_all = build_dfs_from_questions(results)

        # CSV
        csv_bytes = df_all.to_csv(index=False).encode("utf-8")
        csv_key = hashlib.md5(csv_bytes).hexdigest()
        store_result_bytes(csv_key, csv_bytes, "video_questions.csv", "text/csv")

        # Excel
        excel_buf = BytesIO()
        with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
            df_all.to_excel(writer, sheet_name="Questions", index=False)
        excel_buf.seek(0)
        excel_bytes = excel_buf.getvalue()
        excel_key = hashlib.md5(excel_bytes).hexdigest()
        store_result_bytes(excel_key, excel_bytes, "video_questions.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # DOCX
        docx_bytes = build_docx_bytes(results)
        docx_key = hashlib.md5(docx_bytes).hexdigest()
        store_result_bytes(docx_key, docx_bytes, "video_questions.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # counts for this request
        mcq_count_now = len(mcqs)
        desc_count_now = len(descrs)

        # update global state
        STATE["mcq_count"] = STATE.get("mcq_count", 0) + mcq_count_now
        STATE["desc_count"] = STATE.get("desc_count", 0) + desc_count_now

        return {
            "status": "success",
            "mcqCount": mcq_count_now,
            "descCount": desc_count_now,
            "download_keys": {"csv": csv_key, "excel": excel_key, "docx": docx_key},
            "global_state": {
                "pdf_uploads": STATE["pdf_uploads"],
                "last_pdf_pages": STATE["last_pdf_pages"],
                "mcq_count": STATE["mcq_count"],
                "desc_count": STATE["desc_count"]
            },
            "results": results,
            "summary": summary_text,
            "chunks": chunk_summaries
        }
    except Exception as e:
        return {"status": "error", "error": str(e)}





















































































   
        

