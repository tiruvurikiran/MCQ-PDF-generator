# # # descriptive questions and MCQS for PDFs and MCQs for video

import streamlit as st
import fitz  # PyMuPDF
import re
import json
import requests
import random
from typing import List, Dict, Tuple
from io import BytesIO
from docx import Document
import tempfile
import os
import shutil
from faster_whisper import WhisperModel
import cv2
from PIL import Image
import pytesseract
import pandas as pd

import base64
import hashlib
import io
from xlsxwriter import Workbook


# --- Authentication wrapper (place this at top of your script, right after imports) ---


import streamlit_cookies_manager as cookie_manager

# Default users
DEFAULT_USERS = {
	"admin": "password123",
	"user1": "secret"
}

# Function to load users
def load_users():
	try:
		users = dict(st.secrets.get("users") or {})
		if users:
			return users
	except Exception:
		pass
	return DEFAULT_USERS

USERS = load_users()

# Initialize cookies manager
cookies = cookie_manager.CookieManager()

# Wait until cookies are ready
if not cookies.ready():
		st.stop()

# Check if the user is already logged in via cookies
if 'logged_in' not in st.session_state:
	st.session_state['logged_in'] = cookies.get("logged_in", False)
	st.session_state['username'] = cookies.get("username", None)

if "login_failed" not in st.session_state:
	st.session_state["login_failed"] = False

# --- Top-right Logout button (visible only when logged in) ---



# --- Login form (blocks the rest of the script until success) ---
if not st.session_state["logged_in"]:
	st.markdown("### Please sign in to continue")
	u_col, p_col = st.columns(2)
	with u_col:
		username_input = st.text_input("Username", key="login_username")
	with p_col:
		password_input = st.text_input("Password", type="password", key="login_password")

	if st.button("Login"):
		username_try = (username_input or "").strip()
		pwd_try = password_input or ""
		if username_try and USERS.get(username_try) == pwd_try:
			st.session_state["logged_in"] = True
			st.session_state["username"] = username_try
			st.session_state["login_failed"] = False
			cookies["logged_in"] = True
			cookies["username"] = username_try
			cookies.save()
			st.success(f"Welcome, {username_try}!")
			# no explicit rerun needed; Streamlit will rerun automatically after button event
		else:
			st.session_state["login_failed"] = True

	if st.session_state["login_failed"]:
		st.error("Invalid username or password. (Tip: configure users via st.secrets for production.)")

	# prevent the rest of the app from running until logged in
	st.stop()

# ---------------- Dashboard state defaults ----------------
if "pdf_uploads" not in st.session_state:
	st.session_state["pdf_uploads"] = 0  # how many distinct PDF files uploaded
if "last_pdf_hash" not in st.session_state:
	st.session_state["last_pdf_hash"] = None
if "last_pdf_pages" not in st.session_state:
	st.session_state["last_pdf_pages"] = 0
if "mcq_count" not in st.session_state:
	st.session_state["mcq_count"] = 0
if "desc_count" not in st.session_state:
	st.session_state["desc_count"] = 0


# ---------- Custom CSS and Styling ----------
st.markdown(
	"""
	<style>
	/* Hide Streamlit's default header */
	/* Hide Streamlit's footer */
	footer { visibility: hidden; }
	/* Override Streamlit default font */
	html, section, body, [class*="css"]  {
		background: #fff !important;
		max-width: 100%;
	}
	/* Hide hamburger menu and Streamlit branding */
	.st-emotion-cache-1v0mbdj { display: none; }  /* Sidebar collapse icon */
	.st-emotion-cache-13ln4jf { display: none; }  /* Main menu */
	.st-emotion-cache-1d391kg { display: none; }  /* Footer */
	.st-emotion-cache-1w723zb {
		width: 100%;
		padding: 6rem 1rem 10rem;
		max-width: 100%;
	}
	</style>
	""",
	unsafe_allow_html=True
)

# Handle logout click
if "logout" in st.query_params:   # button submitted
	st.session_state["logged_in"] = False
	st.session_state["username"] = None
	st.session_state["login_failed"] = False
	for k in ["pdf_questions", "pdf_docx_bytes", "pdf_generation_meta", "pdf_chaps_select"]:
		st.session_state.pop(k, None)
	st.query_params.clear()
	st.rerun()

# Custom header with logo + logout button (should display at the top when logged in)
encoded_svg = base64.b64encode(open("C:/Users/CITHP/Documents/Summary and MCQS/Logo.svg", "rb").read()).decode()

st.markdown(f"""
	<style>
	header[data-testid="stHeader"] {{
		display: none;
	}}
	header.header {{
		position: fixed;
		top: 0;
		left: 0;
		width: 100%;
		padding: 5px 2rem;
		background: #ffffff;
		box-shadow: -3px 3px 6px rgba(0,0,0,0.16);
		z-index: 99999;
		display: flex;
		align-items: center;
		justify-content: space-between;
	}}
	header.header .logo img {{
		height: 50px;
		width: auto;
	}}
	.spacer-header {{
		height: 70px;
	}}
	.logout-btn button {{
		background-color: #b11226 !important;
		color: white !important;
		border-radius: 5px !important;
		font-size: 14px !important;
		padding: 4px 12px !important;
	}}
	</style>
	<header class="header">
		<div class="logo">
		<img src="data:image/svg+xml;base64,{encoded_svg}" alt="ICFAI Logo">
		</div>
		<div class="logout-btn">
		<form action="#" method="get">
			<button name="logout" type="submit">Logout</button>
		</form>
		</div>
	</header>
	<div class="spacer-header"></div>
""", unsafe_allow_html=True)


st.markdown("""
<style>
/* Remove extra padding Streamlit adds around markdown blocks */
.block-container {
	padding-top: 0rem !important; 
}

/* Remove extra margins around markdown text */
.stMarkdown {
	margin: 0 !important;
	padding: 0 !important;
}
</style>
""", unsafe_allow_html=True)

footer_html = """
<style>
	.footer {
	position: fixed;
	left: 0;
	bottom: 0;
	width: 100%;
	background-color: #b11226;  /* deep red */
	color: white;
	text-align: center;
	padding: 0.75rem 0;
	font-size: 0.9rem;
	z-index: 100;
	}
	/* avoid Streamlit’s bottom bar from overlapping */
	.reportview-container .main footer { visibility: hidden; }
</style>
<div class="footer">
	© 2024 Copyright All Rights Reserved by 
	<strong>The ICFAI Group</strong>.
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)


# Optional video/audio libs
try:
	from moviepy.editor import VideoFileClip
	_HAS_MOVIEPY = True
except Exception:
	_HAS_MOVIEPY = False









st.set_page_config(page_title="PDF & Video → MCQ Generator", layout="wide")
# ------------------ Top dashboard ------------------
dash_col1, dash_col2, dash_col3, dash_col4 = st.columns(4)

with dash_col1:
	st.metric("PDFs uploaded", st.session_state.get("pdf_uploads", 0))

with dash_col2:
	st.metric("Pages (current PDF)", st.session_state.get("last_pdf_pages", 0))

with dash_col3:
	st.metric("MCQs generated", st.session_state.get("mcq_count", 0))

with dash_col4:
	st.metric("Descriptive Qs", st.session_state.get("desc_count", 0))
# ------------------ end dashboard ------------------
st.title("PDF & Video → MCQ Generator")

# --- Inline logout button (place this after st.title(...)) ---

# ------------------- Config -------------------
# Hardcoded Ollama model (hidden)
OLLAMA_MODEL = "llama3"
# Ollama endpoint
OLLAMA_URL = "http://localhost:11434/api/generate"

# ------------------- Utilities -------------------
def clean_text(text: str) -> str:
	"""Remove NULL bytes and control characters for safe DOCX export."""
	if text is None:
		return ""
	return re.sub(r"[\x00-\x1F\x7F]", "", str(text))

def detect_index_range(doc, min_section_hits: int = 3, consecutive_break: int = 2) -> Tuple[int, int]:
	"""
	Heuristic to detect start/end page indices (1-indexed) of a Table of Contents.
	Returns (start_page_1idx, end_page_1idx). Raises ValueError if detection fails.
	"""
	scores = []
	has_contents_flags = []
	for pno in range(doc.page_count):
		try:
			text = doc.load_page(pno).get_text("text") or ""
		except Exception:
			text = ""
		low = text.lower()
		has_contents = bool(re.search(r"\btable of contents\b|\bcontents\b", low))
		count_sections = len(re.findall(r"\b\d{1,2}\.\d+\b", text))
		count_leaders = len(re.findall(r"\.{2,}\s*\d+|\s+\d{1,3}\s*$", text, re.M))
		score = count_sections + 0.6 * count_leaders + (5 if has_contents else 0)
		scores.append(score)
		has_contents_flags.append(has_contents)

	# If explicit 'contents' found, use earliest such page
	if any(has_contents_flags):
		start_idx = next(i for i, f in enumerate(has_contents_flags) if f)
		end_idx = start_idx
		break_count = 0
		for i in range(start_idx + 1, len(scores)):
			if scores[i] >= 1.0:
				end_idx = i
				break_count = 0
			else:
				break_count += 1
				if break_count >= consecutive_break:
					break
		return (start_idx + 1, end_idx + 1)

	# Otherwise, find first page with enough section hits
	start_idx = None
	for i, s in enumerate(scores):
		if s >= min_section_hits:
			start_idx = i
			break
	if start_idx is None:
		raise ValueError("Could not auto-detect contents/index pages. The PDF's TOC format might be unusual.")

	end_idx = start_idx
	gap = 0
	for i in range(start_idx + 1, len(scores)):
		if scores[i] >= 1.0:
			end_idx = i
			gap = 0
		else:
			gap += 1
			if gap >= consecutive_break:
				break

	return (start_idx + 1, end_idx + 1)

# ------------------- Ollama MCQ Generator (with difficulty + explanation) -------------------
def generate_mcqs_ollama(topic: str, context: str = "", full_text: str = "",model: str = OLLAMA_MODEL) -> List[Dict]:
	"""
	Call Ollama and parse Q/A blocks robustly.
	Returns list of dicts: { question, options (list 'A. ...'), answer (single letter or ''), difficulty (optional) }.
	"""
	prompt = f"""
Generate  distinct multiple-choice questions for the topic below. For each question include:
- Exactly 4 labeled options A) B) C) D)
- A single-letter correct answer on its own line: Answer: <A/B/C/D>
- (Optional) Difficulty line: Difficulty: <1-5>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <letter>
Difficulty: <1-5>

Topic: {topic}
Context: {context[:1500]}
"""

	try:
		resp = requests.post(
			OLLAMA_URL,
			json={"model": model, "prompt": prompt, "stream": False, "temperature": 0.3, "max_tokens": 800},
			timeout=240,
		)
		resp.raise_for_status()
		output = resp.json().get("response", "") or ""
	except Exception as e:
		return [{"question": f"Ollama error: {e}", "options": [], "answer": ""}]

	# --- parse robustly ---
	def strip_q_prefix(s: str) -> str:
		# Just clean weird control chars, don’t remove Qn prefix
		return clean_text(s.strip())


	mcqs = []
	# split at Q1., Q2., etc (keep blocks)
	blocks = re.split(r'\n(?=Q\d+\.)', output)
	for block in blocks:
		block = block.strip()
		if not block:
			continue
		lines = [ln.rstrip() for ln in block.splitlines() if ln.strip()]
		if not lines:
			continue

		# find question line (first line that looks like "Qn." or the first non-option line)
		q_line_idx = 0
		for idx, ln in enumerate(lines):
			if re.match(r'^\s*Q\d+\.', ln, re.I):
				q_line_idx = idx
				break
			# if first non-option line and not an intro, take it
			if not re.match(r'^[A-D][\)\.\-:]', ln, re.I) and not re.search(r'(here are|multiple[-\s]?choice|based on the topic)', ln, re.I):
				q_line_idx = idx
				break

		q_line = strip_q_prefix(lines[q_line_idx])

		# collect options starting immediately after q_line_idx
		opts = []
		opt_end_idx = q_line_idx
		for j in range(q_line_idx + 1, len(lines)):
			m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
			if m:
				label = m.group(1).upper()
				text = m.group(2).strip()
				opts.append(f"{label}. {text}")
				opt_end_idx = j
			else:
				# stop collecting options when a non-option line encountered
				break

		# look for Answer: in the next few lines after options (opt_end_idx + 1.. +4)
		answer = ""
		difficulty = ""
		look_start = opt_end_idx + 1
		look_end = min(len(lines), opt_end_idx + 6)
		for k in range(look_start, look_end):
			ln = lines[k]
			# answer line like "Answer: B" or "Correct: B)"
			m_ans = re.search(r'(?i)\b(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', ln)
			if m_ans:
				answer = m_ans.group(1).upper()
				continue
			# difficulty line like "Difficulty: 3" or "Level: 3"
			m_diff = re.search(r'(?i)\b(?:difficulty|level)[:\s\-]*\(?\s*([1-5])\s*\)?', ln)
			if m_diff:
				difficulty = m_diff.group(1)
				continue
			# sometimes the model puts a single-letter line like "B" or "B)"
			m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', ln, re.I)
			if m_single and not answer:
				answer = m_single.group(1).upper()

		# As a last resort, if no explicit answer found, look for "Answer:" anywhere in block:
		if not answer:
			m_any = re.search(r'(?i)\banswer[:\s\-]*\(?\s*([A-D])\s*\)?', block)
			if m_any:
				answer = m_any.group(1).upper()

		# Accept only if question present and at least 2 options
		if q_line and len(opts) >= 2:
			mcqs.append({
				"question": clean_text(q_line),
				"options": [clean_text(o) for o in opts],
				"answer": clean_text(answer),
				"difficulty": clean_text(difficulty)
			})

	return mcqs


# ------------------- Descriptive questions (Ollama) -------------------
def parse_descriptive_freeform(output: str) -> List[str]:
	"""
	Parse descriptive/open-ended questions returned as:
	Q1. <question text>
	Q2. <question text>
	...
	Returns list of question strings.
	"""
	items: List[str] = []
	if not output:
		return items
	# Split by Qn. (keep blocks)
	blocks = re.split(r'\n(?=Q\d+\.)', output)
	for block in blocks:
		block = block.strip()
		if not block:
			continue
		# first line containing Qn.
		first_line = block.splitlines()[0].strip()
		m = re.match(r'^\s*Q\d+\.\s*(.*)$', first_line, re.I)
		if m:
			q = m.group(1).strip()
		else:
			# fallback: take whole block as a single question
			q = block.strip()
		if q:
			items.append(q)
	return items

def generate_descriptive_with_answers(topic: str, context: str = "", model: str = OLLAMA_MODEL, num_qs: int = 3) -> List[Dict]:
		"""
		Generate descriptive questions WITH answers and difficulty levels.
		Returns list of dicts: { question, answer, difficulty }.
		"""
		prompt = f"""
Generate {num_qs} descriptive / short-answer / essay-style questions for the topic below.
For each question, also provide:
- Correct answer
- Difficulty level (1-5)

Return exactly in this format:

Q1. <question text>
Answer: <answer text>
Difficulty: <1-5>

Do not add extra commentary.

Topic: {topic}
Context: {context[:1500]}
"""
		try:
				resp = requests.post(
						OLLAMA_URL,
						json={"model": model, "prompt": prompt, "stream": False, "temperature": 0.3, "max_tokens": 600},
						timeout=240,
				)
				resp.raise_for_status()
				out = resp.json().get("response", "") or ""
		except Exception as e:
				return [{"question": f"Ollama error: {e}", "answer": "", "difficulty": ""}]

		# Parse Q/Answer/Difficulty blocks
		blocks = re.split(r'\n(?=Q\d+\.)', out)
		results = []
		for block in blocks:
				block = block.strip()
				if not block:
						continue
				lines = block.splitlines()
				question = ""
				answer = ""
				difficulty = ""
				for ln in lines:
						ln = ln.strip()
						if ln.lower().startswith("q"):
								question = re.sub(r'^q\d+\.\s*', '', ln, flags=re.I).strip()
						elif ln.lower().startswith("answer:"):
								answer = ln.split(":", 1)[1].strip()
						elif ln.lower().startswith("difficulty:"):
								difficulty = ln.split(":", 1)[1].strip()
				if question:
						results.append({"question": question, "answer": answer, "difficulty": difficulty})
		return results



# ---------- Descriptive rendering: split, clean and display ----------
def strip_q_prefix(s: str) -> str:
	"""Remove leading Qn./Question n: prefixes."""
	return re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', (s or ""), flags=re.I).strip()

def is_intro_line(s: str) -> bool:
	"""Detect obvious intro/header lines to drop (like 'Here are...' or 'Following questions...')."""
	return bool(re.search(r'^(here are|the following|following questions|below are|here\'s|here is)\b', s.strip(), re.I))

def split_descriptive_block(s: str) -> List[str]:
	"""
	Normalize a single descriptive string:
	- If it contains Qn. markers, split on them.
	- Otherwise return the cleaned string (unless it's an intro header).
	"""
	if not s or not s.strip():
		return []
	# If contains explicit Qn markers, split at each 'Q<number>.' (case-insensitive).
	if re.search(r'\bq\d+\.', s, re.I):
		parts = re.split(r'(?i)(?=\bq\d+\.)', s)  # keep the Qn. at the start of each part
		out = []
		for p in parts:
			p_clean = strip_q_prefix(p)
			if not p_clean:
				continue
			if is_intro_line(p_clean):
				# skip pure intro fragments
				continue
			out.append(p_clean)
		return out
	else:
		# no explicit Qn found — just return the cleaned string unless it's clearly an intro header
		candidate = strip_q_prefix(s)
		if not candidate:
			return []
		if is_intro_line(candidate):
			return []
		return [candidate]


# ------------------- PDF Flow -------------------
# ------------------- PDF Flow with Descriptive Answers -------------------
st.header("PDF → Topics → MCQs & Descriptive")
st.markdown("""
<style>
/* Shrink the file uploader box */
[data-testid="stFileUploader"] section {
		padding: 0.25rem !important;
		font-size: 0.85rem !important;
		min-height: 2.2rem !important;
}
[data-testid="stFileUploader"] button {
		padding: 0.2rem 0.6rem !important;
		font-size: 0.8rem !important;
		height: auto !important;
}
[data-testid="stFileUploader"] {
		max-width: 600px !important;
		margin: 0 !important;
}
</style>
""", unsafe_allow_html=True)

pdf_file = st.file_uploader("Upload a PDF (with Table of Contents)", type=["pdf"], key="pdf_uploader")

if pdf_file is not None:
		try:
				pdf_bytes = pdf_file.read()
				pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
				md5 = hashlib.md5(pdf_bytes).hexdigest()
				if st.session_state.get("last_pdf_hash") != md5:
						st.session_state["pdf_uploads"] = st.session_state.get("pdf_uploads", 0) + 1
						st.session_state["last_pdf_hash"] = md5
				st.session_state["last_pdf_pages"] = getattr(pdf_doc, "page_count", 0)
		except Exception as e:
				st.error(f"Failed to read PDF: {e}")
				pdf_doc = None

		if pdf_doc:
				# Auto-detect contents pages
				try:
						detected_start, detected_end = detect_index_range(pdf_doc)
						st.success(f"Auto-detected contents pages: {detected_start} — {detected_end}  "
											 f"📖 Textbook pages: {pdf_doc.page_count}")
				except Exception as e:
						st.error(f"Auto-detection failed: {e}")
						detected_start = detected_end = None

				if detected_start and detected_end:
						# Extract TOC text
						index_text = "\n".join([pdf_doc.load_page(p).get_text("text") or "" 
																		for p in range(detected_start - 1, detected_end)])
						# Extract full PDF text
						full_text = "\n".join([pdf_doc.load_page(p).get_text("text") or "" 
																	 for p in range(pdf_doc.page_count)])

						# Parse headings from TOC
						raw_matches = re.findall(r"(\d{1,2}\.\d+)\s+(.+?)\s+(\d{1,4})\b", index_text)
						matches = []
						if raw_matches:
								for num, title, pno in raw_matches:
										title_clean = re.sub(r"\.{2,}|\.{3,}", ".", title)
										title_clean = re.sub(r"\.{2,}.*$", "", title_clean)
										title_clean = title_clean.strip(' .\t')
										title_clean = clean_text(title_clean)
										try:
												page_no = int(pno)
										except Exception:
												page_no = None
										if title_clean:
												matches.append((num.strip(), title_clean, page_no))
						else:
								raw_matches2 = re.findall(r"(\d{1,2}\.\d+)\s+([^\n\r]+)", index_text)
								for num, title in raw_matches2:
										title_clean = re.sub(r"\.{2,}|\.{3,}", ".", title)
										title_clean = re.sub(r"\.{2,}.*$", "", title_clean)
										title_clean = re.sub(r"\s+\d+$", "", title_clean)
										title_clean = title_clean.strip(' .\t')
										title_clean = clean_text(title_clean)
										if title_clean:
												matches.append((num.strip(), title_clean, None))

						if not matches:
								st.warning("No numeric headings found on the detected contents pages.")
						else:
								# Group by chapter
								chapters = {}
								for subnum, title, page_no in matches:
										chap_no = int(subnum.split('.')[0]) if subnum.split('.')[0].isdigit() else 0
										chapters.setdefault(chap_no, []).append({'subnum': subnum, 'title': title, 'page': page_no})

								st.write(f"Found {len(matches)} subheadings across {len(chapters)} chapters.")
								with st.expander("Preview parsed headings (from contents)"):
										for chap_no in sorted(chapters.keys()):
												st.write(f"Chapter {chap_no} — {len(chapters[chap_no])} topics")
												for t in chapters[chap_no]:
														st.write(f"  {t['subnum']} — {t['title']} (page: {t.get('page')})")

								# User scope & type selections
								col1, col2, col3 = st.columns([1, 1.2, 1.2])
								with col1:
										generate_scope = st.radio("Generate questions for:", ("All topics", "Selected chapters only"), key="pdf_scope")
								with col2:
										mcq_source = st.radio("MCQ source:", ("Llama3 (open-ended)", "From textbook (use chapter content)"),
																					index=0, key="mcq_source")
								with col3:
										question_type = st.radio("Question type:", ("MCQs", "Descriptive", "Both"), index=0, key="question_type")

								selected_chapters = []
								if generate_scope == "Selected chapters only":
										options = [f"Chapter {c}: {len(chapters[c])} topics" for c in sorted(chapters.keys())]
										if "pdf_chaps_select" not in st.session_state:
												st.session_state["pdf_chaps_select"] = []
										chosen = st.multiselect("Choose chapters", options, default=st.session_state.get("pdf_chaps_select", []), key="pdf_chaps_select")
										chosen = [c for c in (st.session_state.get("pdf_chaps_select", []) or []) if c in options]
										for s in chosen:
												try:
														selected_chapters.append(int(s.split()[1].strip(':')))
												except Exception:
														pass

								# Ensure session_state keys
								for k in ['pdf_questions', 'pdf_docx_bytes', 'pdf_generation_meta']:
										if k not in st.session_state:
												st.session_state[k] = {} if k=='pdf_questions' else None

								# Check if questions have been generated before
								if "pdf_questions" not in st.session_state:
										st.session_state["pdf_questions"] = {}

								# Check if the DOCX content is already available
								if "pdf_docx_bytes" not in st.session_state:
										st.session_state["pdf_docx_bytes"] = None

								# Generate button to generate MCQs and descriptive questions
								if st.button("Generate MCQs and descriptive Questions"):
										st.session_state['pdf_questions'] = {}
										st.session_state['pdf_docx_bytes'] = None
										st.session_state['pdf_generation_meta'] = {}

										topics = []
										# Use a variable to determine which chapters to include in the topics
										if generate_scope == "All topics":
												for chap_no in sorted(chapters.keys()):
														for t in chapters[chap_no]:
																topics.append({'chapter': chap_no, 'subnum': t['subnum'], 'title': t['title'], 'page': t.get('page')})
										else:
												for chap_no in selected_chapters:
														for t in chapters.get(chap_no, []):
																topics.append({'chapter': chap_no, 'subnum': t['subnum'], 'title': t['title'], 'page': t.get('page')})

										results_tmp = {}
										docx_pdf = Document()
										docx_pdf.add_heading("Generated Questions", level=1)
										topic_placeholders = []

										for idx_topic, topic in enumerate(topics, start=1):
												topic_title = topic['title']
												topic_ph = st.empty()
												topic_placeholders.append(topic_ph)
												with topic_ph.container():
														st.subheader(f" {topic_title}")

												# Build context for MCQ generation
												if mcq_source == "From textbook (use chapter content)":
														found_idx = full_text.lower().find(topic_title.lower())
														if found_idx != -1:
																start = max(0, found_idx - 1500)
																end = min(len(full_text), found_idx + 3500)
																context = full_text[start:end]
														else:
																pg = topic.get('page')
																if pg and isinstance(pg, int) and 1 <= pg <= pdf_doc.page_count:
																		part = [pdf_doc.load_page(p).get_text("text") or "" for p in range(max(0, pg-1), min(pdf_doc.page_count, pg+2))]
																		context = "\n".join(part)
																else:
																		context = index_text
												else:
														context = ""

												results_tmp[topic_title] = {"mcqs": [], "descriptive": []}

												# Generate questions
												try:
														with st.spinner(f"Generating questions for: {topic_title}"):
																mcqs_local = []
																descrs_local = []

																if question_type in ("MCQs", "Both"):
																		mcqs_local = generate_mcqs_ollama(topic_title, context=context, full_text=index_text, model=OLLAMA_MODEL) or []
																		results_tmp[topic_title]["mcqs"] = mcqs_local

																if question_type in ("Descriptive", "Both"):
																		descrs_local = generate_descriptive_with_answers(topic_title, context=context, model=OLLAMA_MODEL, num_qs=3) or []
																		results_tmp[topic_title]["descriptive"] = descrs_local

																# Render per-topic output
																with topic_ph.container():
																		st.subheader(f"{topic_title}")
																		# MCQs
																		for mcq in mcqs_local:
																				st.markdown(f"**{mcq.get('question', '')}**  —  _Level {mcq.get('difficulty', 'N/A')}_")
																				for opt in mcq.get("options", []):
																						st.write(f"- {opt}")
																				if mcq.get("answer"):
																						st.success(f"✅ Answer: {mcq['answer']}")

																		# Descriptive
																		if descrs_local:
																				st.write("**Descriptive / Open-ended questions:**")
																				for i, dq in enumerate(descrs_local, start=1):
																						st.markdown(f"- Q{i}. {dq.get('question', '')}  —  _Level {dq.get('difficulty', 'N/A')}_")
																						if dq.get("answer"):
																								st.success(f"✅ Answer: {dq['answer']}")

												except Exception as e:
														st.error(f"Error generating questions for {topic_title}: {e}")

										# Save the results in session state
										# After you finish generating 'results_tmp' and 'docx_pdf' (inside the Generate button block)
										# ---------------- Save generated results into session_state (make bytes persistent) ----------------

										def build_docx_bytes(questions_data: dict) -> bytes:
											"""
											Build a DOCX file from questions_data (MCQs + Descriptive).
											Always include both if present.
											"""
											doc = Document()
											doc.add_heading("Generated Questions", level=1)

											for topic_title, blocks in questions_data.items():
													doc.add_heading(topic_title, level=2)

													# MCQs
													mcqs = blocks.get("mcqs", []) or []
													if mcqs:
															doc.add_paragraph("Multiple Choice Questions:", style="List Bullet")
															for idx, mcq in enumerate(mcqs, start=1):
																	qtext = mcq.get("question", "")
																	doc.add_paragraph(f"{idx}. {qtext}")
																	opts = mcq.get("options", []) or []
																	for opt in opts:
																			doc.add_paragraph(f"    {opt}")
																	answer = mcq.get("answer", "")
																	diff = mcq.get("difficulty", "N/A")
																	if answer:
																			doc.add_paragraph(f"    Answer: {answer}    Difficulty: {diff}")
																	else:
																			doc.add_paragraph(f"    Difficulty: {diff}")
																	doc.add_paragraph("")

													# Descriptive
													descrs = blocks.get("descriptive", []) or []
													if descrs:
															doc.add_paragraph("Descriptive / Short-answer Questions:", style="List Bullet")
															for idx, dq in enumerate(descrs, start=1):
																	if isinstance(dq, dict):
																			q = dq.get("question", "")
																			a = dq.get("answer", "")
																			diff = dq.get("difficulty", "N/A")
																	else:
																			q = str(dq)
																			a, diff = "", "N/A"
																	doc.add_paragraph(f"{idx}. {q}")
																	if a:
																			doc.add_paragraph(f"    Answer: {a}")
																	doc.add_paragraph(f"    Difficulty: {diff}")
																	doc.add_paragraph("")

											buf = BytesIO()
											doc.save(buf)
											buf.seek(0)
											return buf.getvalue()


										st.session_state['pdf_questions'] = results_tmp  # store questions dict

										# Save DOCX bytes (you already did this but ensure it's bytes in session_state)
										try:
												st.session_state['pdf_docx_bytes'] = build_docx_bytes(results_tmp)
										except Exception as e:
												st.error(f"Failed to build DOCX: {e}")
												st.session_state['pdf_docx_bytes'] = None

										# Create CSV bytes and Excel bytes from the same session data and store them
										def build_dfs_from_questions(questions_data: dict):
												"""Return (df: pandas.DataFrame) combining MCQs + Descriptive for CSV/Excel."""
												rows = []
												for topic_title, topic_data in questions_data.items():
														# MCQs
														for mcq in topic_data.get("mcqs", []):
																# normalize options into separate columns
																opts = mcq.get("options") or []
																rows.append({
																		"Topic": topic_title,
																		"Type": "MCQ",
																		"Question": mcq.get("question", ""),
																		"Option A": opts[0] if len(opts) > 0 else "",
																		"Option B": opts[1] if len(opts) > 1 else "",
																		"Option C": opts[2] if len(opts) > 2 else "",
																		"Option D": opts[3] if len(opts) > 3 else "",
																		"Answer": mcq.get("answer", ""),
																		"Difficulty": mcq.get("difficulty", "N/A"),
																		"Descriptive Answer": ""
																})
														# Descriptive (already contains answer + difficulty in your version)
														for dq in topic_data.get("descriptive", []):
																rows.append({
																		"Topic": topic_title,
																		"Type": "Descriptive",
																		"Question": dq.get("question", ""),
																		"Option A": "", "Option B": "", "Option C": "", "Option D": "",
																		"Answer": "",  # MCQ answer empty
																		"Difficulty": dq.get("difficulty", "N/A"),
																		"Descriptive Answer": dq.get("answer", "")
																})
												df = pd.DataFrame(rows)
												return df

										# Build dataframe and save CSV/Excel bytes
										try:
												df_all = build_dfs_from_questions(results_tmp)
												# CSV bytes
												csv_bytes = df_all.to_csv(index=False).encode("utf-8")
												st.session_state['pdf_csv_bytes'] = csv_bytes

												# Excel bytes
												excel_buf = BytesIO()
												with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
														df_all.to_excel(writer, sheet_name="Questions", index=False)
												excel_buf.seek(0)
												st.session_state['pdf_excel_bytes'] = excel_buf.getvalue()
										except Exception as e:
												st.error(f"Failed to build CSV/Excel: {e}")
												st.session_state['pdf_csv_bytes'] = None
												st.session_state['pdf_excel_bytes'] = None

										# update counts
										st.session_state["mcq_count"] = sum(len(t.get("mcqs", [])) for t in results_tmp.values())
										st.session_state["desc_count"] = sum(len(t.get("descriptive", [])) for t in results_tmp.values())

										# Force a rerun so top dashboard updates immediately
										st.rerun()

																		
								 # Always render saved results from session_state (so they persist after rerun and after downloads)
								if st.session_state.get('pdf_questions'):
										st.markdown("### Generated questions (saved)")
										for topic_title, blocks in st.session_state['pdf_questions'].items():
												st.subheader(topic_title)   # Topic name shown above questions

												# MCQs
												for mcq in blocks.get("mcqs", []) or []:
														q = mcq.get("question") or ""
														diff = mcq.get("difficulty") or "N/A"
														st.markdown(f"**{q}**  —  _Level {diff}_")
														for opt in (mcq.get("options") or []):
																if opt:
																		st.write(f"- {opt}")
														if mcq.get("answer"):
																st.success(f"✅ Correct: {mcq.get('answer')}")
														else:
																st.info("✅ Correct: ")

												# Descriptive with answers + difficulty
												raw_descrs = blocks.get("descriptive", []) or []
												if raw_descrs:
														st.write("**Descriptive / Open-ended questions:**")
														for i, dq in enumerate(raw_descrs, start=1):
																# if you stored descriptive as dict with question/answer/difficulty:
																q_text = dq.get("question", "") if isinstance(dq, dict) else dq
																ans_text = dq.get("answer", "") if isinstance(dq, dict) else ""
																difficulty = dq.get("difficulty", "N/A") if isinstance(dq, dict) else "N/A"
																st.markdown(f"- Q{i}. {q_text}  —  _Level {difficulty}_")
																if ans_text:
																		st.success(f"✅ Answer: {ans_text}")

										# Download buttons using bytes preserved in session_state (these won't disappear UI)
										cols = st.columns(3)
										with cols[0]:
												if st.session_state.get('pdf_docx_bytes'):
														st.download_button(
																label="Download DOCX",
																data=st.session_state['pdf_docx_bytes'],
																file_name="generated_questions.docx",
																mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
																key="download_docx"
														)
										with cols[1]:
												if st.session_state.get('pdf_excel_bytes'):
														st.download_button(
																label="Download Excel",
																data=st.session_state['pdf_excel_bytes'],
																file_name="generated_questions.xlsx",
																mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
																key="download_excel"
														)
										with cols[2]:
												if st.session_state.get('pdf_csv_bytes'):
														st.download_button(
																label="Download CSV",
																data=st.session_state['pdf_csv_bytes'],
																file_name="generated_questions.csv",
																mime="text/csv",
																key="download_csv"
														)




				
										


				
				
 


 #  #  # MCQs for video using summary 

# summary using Bart, nltk 
import streamlit as st
import tempfile, os, re, io
from moviepy.editor import VideoFileClip
from collections import Counter
from typing import List, Dict
import requests
from docx import Document

from transformers import pipeline
import nltk
from nltk.tokenize import sent_tokenize
nltk.download('punkt', quiet=True)

CHUNK_WORDS = 800
SUMMARY_MAX_LENGTH = 120
SUMMARY_MIN_LENGTH = 30
SUMMARIZER_MODEL = "facebook/bart-large-cnn"

def split_transcript_into_chunks_by_words(transcript: str, chunk_words: int = CHUNK_WORDS) -> List[str]:
		sentences = sent_tokenize(transcript)
		chunks, current, current_words = [], [], 0
		for s in sentences:
				wcount = len(s.split())
				if current_words + wcount > chunk_words and current:
						chunks.append(" ".join(current))
						current, current_words = [s], wcount
				else:
						current.append(s)
						current_words += wcount
		if current:
				chunks.append(" ".join(current))
		return chunks

def summarizer_pipeline(model_name=SUMMARIZER_MODEL):
		return pipeline("summarization", model=model_name, device=-1)  # CPU

def summarize_chunks(chunks: List[str], summarizer) -> List[str]:
		summaries = []
		for c in chunks:
				try:
						out = summarizer(c, max_length=400, min_length=100, do_sample=False)
						summary_text = out[0]['summary_text'].strip()
				except Exception:
						summary_text = " ".join(c.split()[:SUMMARY_MIN_LENGTH])
				summaries.append(summary_text)
		return summaries

def combine_and_summarize_summaries(summaries: List[str]) -> str:
		"""
		Return an overall summary that is exactly the concatenation of chunk summaries.
		Keeps chunk order and separators intact.
		"""
		if not summaries:
				return ""
		# join with a blank line between chunk summaries so readability is preserved
		overall = "\n\n".join(summaries)
		return overall


def summarize_transcript_with_bart(transcript: str) -> dict:
		"""
		Splits transcript into chunks and summarizes each chunk (via BART).
		The overall summary is the exact concatenation of the chunk summaries (no further condensation).
		"""
		if not transcript or not transcript.strip():
				return {"overall": "", "chunks": []}

		chunks = split_transcript_into_chunks_by_words(transcript, CHUNK_WORDS)
		summarizer = summarizer_pipeline(SUMMARIZER_MODEL)
		chunk_summaries = summarize_chunks(chunks, summarizer)

		overall_summary = combine_and_summarize_summaries(chunk_summaries)
		return {"overall": overall_summary, "chunks": chunk_summaries}



# ---------------- optional libs ----------------
try:
	import whisper
	_HAS_WHISPER = True
except Exception:
	_HAS_WHISPER = False

# ---------------- helpers ----------------
STOPWORDS_SUMMARY = {
	"the","and","for","that","with","this","from","which","these","those",
	"using","use","used","also","etc","about","they","their","will","can","we","you",
	"a","an","in","on","of","to","is","are","was","were","be","it","its"
}

def summarize_text(transcript: str, model: str = "llama3", max_words: int = 200) -> str:
	"""
	Use Ollama to generate an abstractive summary of the transcript.
	"""
	prompt = f"""
	Summarize the following transcript into a clear, concise overall summary (max {max_words} words).
	Focus only on the main ideas, not minor details.

	Transcript:
	{transcript}
	"""

	try:
		resp = requests.post(
			OLLAMA_URL,
			json={"model": model, "prompt": prompt, "stream": False, "temperature": 0.3, "max_tokens": 400},
			timeout=600
		)
		resp.raise_for_status()
		summary = resp.json().get("response", "").strip()
		return summary
	except Exception as e:
		return f"Summary generation error: {e}"


# Robust MCQ parser (accepts many model output formats)
def parse_mcqs_freeform(output: str) -> List[Dict]:
	mcqs = []
	if not output:
		return mcqs
	raw_lines = [ln.rstrip() for ln in output.splitlines() if ln.strip()]
	# drop very generic intro / header-only lines
	lines = []
	for ln in raw_lines:
		if re.search(r"(here are|multiple[-\s]?choice questions|based on the summary|based on the topic|following questions|the following)", ln, re.I):
			continue
		if re.match(r'^\s*(?:question|q)\s*\d+\b[:.\s-]*$', ln, re.I):
			continue
		lines.append(ln.strip())

	i = 0
	while i < len(lines):
		ln = lines[i]
		# skip stray option lines until we find a question
		if re.match(r'^[A-D][\)\.\-:]\s+', ln, re.I):
			i += 1
			continue
		question_text = re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', ln, flags=re.I).strip()
		if len(question_text) < 3:
			i += 1
			continue
		# collect options
		opts = []
		opt_map = {}
		j = i + 1
		while j < len(lines) and len(opts) < 4:
			if re.match(r'^[A-D][\)\.\-:]\s+', lines[j], re.I):
				m = re.match(r'^([A-D])[\)\.\-:]\s*(.*)$', lines[j], re.I)
				if m:
					label = m.group(1).upper()
					text = m.group(2).strip()
					formatted = f"{label}. {text}"
					opts.append(formatted)
					opt_map[label] = formatted
				else:
					opts.append(lines[j].strip())
				j += 1
			else:
				break
		# look ahead for Answer:
		answer = ""
		look_end = min(len(lines), j + 6)
		for k in range(j, look_end):
			candidate = lines[k].strip()
			m_ans = re.match(r'(?i)^\s*(?:answer|correct)[:\s\-]*\(?\s*([A-D])\s*\)?', candidate)
			if m_ans:
				answer = m_ans.group(1).upper()
				break
			m_single = re.match(r'^\s*([A-D])[\)\.\s]*$', candidate, re.I)
			if m_single:
				answer = m_single.group(1).upper()
				break
		if answer and answer not in opt_map:
			answer = ""  # validate
		if question_text and len(opts) >= 2:
			mcqs.append({"question": question_text, "options": opts, "answer": answer})
		i = j if j > i else i + 1
	return mcqs

# ---------------- transcription ----------------
def split_audio(audio_path: str, chunk_length_sec: int = 300) -> List[str]:
	try:
		from pydub import AudioSegment
	except Exception:
		return [audio_path]
	import wave, contextlib
	with contextlib.closing(wave.open(audio_path, 'rb')) as wf:
		rate = wf.getframerate()
		n_frames = wf.getnframes()
		total_sec = n_frames / float(rate)
	if total_sec <= chunk_length_sec:
		return [audio_path]
	audio = AudioSegment.from_wav(audio_path)
	chunk_files = []
	for start_ms in range(0, len(audio), chunk_length_sec * 1000):
		chunk = audio[start_ms:start_ms + chunk_length_sec * 1000]
		tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
		chunk.export(tmp.name, format="wav")
		chunk_files.append(tmp.name)
	return chunk_files

def transcribe_video_bytes(video_bytes: bytes, whisper_model_name: str = "small") -> str:
	if not _HAS_WHISPER:
		raise RuntimeError("Whisper is not installed. pip install openai-whisper")
	# write video
	vf = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
	vf.write(video_bytes); vf.flush(); vf.close()
	audio_path = None
	try:
		clip = VideoFileClip(vf.name)
		af = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
		audio_path = af.name
		clip.audio.write_audiofile(audio_path, logger=None)
		clip.close()
		chunk_files = split_audio(audio_path)
		model = whisper.load_model(whisper_model_name)
		full_text = ""
		for c in chunk_files:
			res = model.transcribe(c)
			text = res.get("text", "").strip()
			if text:
				full_text += text + " "
			try:
				if c != audio_path and os.path.exists(c):
					os.remove(c)
			except Exception:
				pass
		return full_text.strip()
	finally:
		try:
			if os.path.exists(vf.name): os.remove(vf.name)
		except Exception:
			pass
		try:
			if audio_path and os.path.exists(audio_path): os.remove(audio_path)
		except Exception:
			pass

# ---------------- Ollama MCQ generator ----------------
OLLAMA_URL = "http://localhost:11434/api/generate"  # adjust if needed
def generate_mcqs_from_summary(summary: str, num_qs: int = 5, model: str = "llama3") -> List[Dict]:
	# ask the model to generate `num_qs` covering the summary
	prompt = f"""
Generate 10 distinct multiple-choice questions that cover the following summary.
For each question include:
- Exactly 4 labeled options A) B) C) D) 
- A single-letter answer line like: Answer: <A/B/C/D>

Use exactly this format; do not add extra commentary or code fences.

Q1. <question text>
A) <option A>
B) <option B>
C) <option C>
D) <option D>
Answer: <A/B/C/D>

Summary:
{summary}


"""
	try:
		resp = requests.post(OLLAMA_URL, json={"model": model, "prompt": prompt, "stream": False, "temperature": 0.3, "max_tokens": 800}, timeout=600)
		resp.raise_for_status()
		output = resp.json().get("response","") or ""
	except Exception as e:
		return [{"question": f"Ollama error: {e}", "options": [], "answer": ""}]
	return parse_mcqs_freeform(output)

# ---------------- Streamlit app ----------------
st.set_page_config(page_title="Video → Summary → MCQs", layout="wide")
st.title("Video → Summary → MCQs")

# session-state
if "transcript" not in st.session_state: st.session_state.transcript = ""
if "summary" not in st.session_state: st.session_state.summary = ""
if "mcqs" not in st.session_state: st.session_state.mcqs = []

video_file = st.file_uploader("Upload video (mp4/mov/mkv/avi)", type=["mp4","mov","mkv","avi"], key="vid_file")

if video_file is not None:
	if st.button("Transcribe video", key="transcribe_btn"):
		if not _HAS_WHISPER:
			st.error("Whisper not available. Install openai-whisper.")
		else:
			with st.spinner("Transcribing video (may take a while)..."):
				try:
					st.session_state.transcript = transcribe_video_bytes(video_file.read(), whisper_model_name="small")
					st.success("Transcription complete.")
					# auto-generate summary
					summaries = summarize_transcript_with_bart(st.session_state.transcript)
					st.session_state.summary = summaries["overall"]
					st.session_state.chunk_summaries = summaries["chunks"]

					st.session_state.mcqs = []
				except Exception as e:
					st.error(f"Transcription failed: {e}")
					st.session_state.transcript = ""
					st.session_state.summary = ""
					st.session_state.mcqs = []

if st.session_state.transcript:
	with st.expander("Transcript (preview)"):
		st.text_area("transcript_area", st.session_state.transcript, height=300, key="transcript_area")

	st.subheader("Summary (auto-generated)")
	st.write(st.session_state.summary)
	if st.session_state.get("chunk_summaries"):
			with st.expander("Detailed chunk summaries"):
					for i, chunk_sum in enumerate(st.session_state.chunk_summaries, start=1):
							st.markdown(f"**Chunk {i}:** {chunk_sum}")

	st.markdown("---")
	st.subheader("Generate MCQs from summary")
	NUM_QS_DEFAULT = 10
	MODEL_DEFAULT = "llama3"

	if st.button("Generate MCQs from summary", key="gen_mcqs_summary"):
		if not st.session_state.summary:
			st.warning("No summary available.")
		else:
			with st.spinner("Generating MCQs from summary..."):
				st.session_state.mcqs = generate_mcqs_from_summary(st.session_state.summary, num_qs=NUM_QS_DEFAULT, model=MODEL_DEFAULT)
			st.success("MCQs generated.")

	# display generated MCQs
	if st.session_state.mcqs:
		st.subheader("Generated MCQs")
		for idx, mcq in enumerate(st.session_state.mcqs, start=1):
			qtext = re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', mcq.get("question",""), flags=re.I).strip()
			st.markdown(f"**Q{idx}. {qtext}**")
			for opt in mcq.get("options", []):
				st.write(f"- {opt}")
			ans = mcq.get("answer") or "Not specified"
			if ans != "Not specified":
				st.success(f"✅ Answer: {ans}")
			else:
				st.info(f"✅ Answer: {ans}")
			st.write("---")

		# download as DOCX
		if st.session_state.mcqs:
			doc = Document()
			doc.add_heading("MCQs generated from video summary", level=1)
			doc.add_paragraph("Summary:")
			doc.add_paragraph(st.session_state.summary)
			doc.add_paragraph("")
			doc.add_heading("Questions", level=2)
			for idx, mcq in enumerate(st.session_state.mcqs, start=1):
				qtext = re.sub(r'^\s*(?:q|question)\s*\d+\s*[:.\-\)]*\s*', '', mcq.get("question",""), flags=re.I).strip()
				doc.add_paragraph(f"Q{idx}. {qtext}")
				for opt in mcq.get("options", []):
					doc.add_paragraph(opt, style="List Bullet")
				ans = mcq.get("answer") or "Not specified"
				doc.add_paragraph(f"Answer: {ans}")
				doc.add_paragraph("")
			buf = io.BytesIO()
			doc.save(buf)
			buf.seek(0)
			st.download_button(
				"Download MCQs (DOCX)",
				data=buf,
				file_name="video_summary_mcqs.docx",
				mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
				key="download_mcqs"
			)













 



# # # PDF text to video generation for multiple languages 

import os
import re
import textwrap
import requests
import numpy as np
import streamlit as st
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfReader
from gtts import gTTS
from moviepy.editor import AudioFileClip, concatenate_videoclips, ImageClip, VideoFileClip, CompositeVideoClip
from collections import Counter
import concurrent.futures
import functools
from textwrap import shorten
import pysrt
from moviepy.video.tools.subtitles import SubtitlesClip
import os, textwrap, numpy as np
import pyttsx3
import time


from googletrans import Translator
translator = Translator()



import streamlit as st




# Custom CSS to override Streamlit's default styles
st.markdown(
	"""
	<style>
	/* Hide Streamlit's default header */

	/* Hide Streamlit's footer */
	footer { visibility: hidden; }

	/* Override Streamlit default font */
	html, section, body, [class*="css"]  {
		
		background: #fff !important;
		 max-width: 100%;
	}

	/* Hide hamburger menu and Streamlit branding */
	.st-emotion-cache-1v0mbdj { display: none; }  /* Sidebar collapse icon */
	.st-emotion-cache-13ln4jf { display: none; }  /* Main menu */
	.st-emotion-cache-1d391kg { display: none; }  /* Footer */


	.st-emotion-cache-1w723zb {
	width: 100%;
	padding: 6rem 1rem 10rem;
	max-width: 100%;
}

	</style>
	""",
	unsafe_allow_html=True
)






# ─── CONFIG ───────────────────────────────────────────────────────────
TARGET_WORDS = 1800
AUDIO_DIR = "audio"
SRT_DIR   = "srt"
VIDEO_DIR = "videos"
IMG_DIR   = "images"
UNSPLASH_ACCESS_KEY = os.getenv("UNSPLASH_ACCESS_KEY", "")
for d in (AUDIO_DIR, SRT_DIR, VIDEO_DIR, IMG_DIR):
	os.makedirs(d, exist_ok=True)







# Hide Streamlit menu and footer
hide_menu = """
	<style>
		#MainMenu {visibility: hidden;}
		footer {visibility: hidden;}
		.stDeployButton {visibility: hidden;}
		[data-testid="stToolbar"] {visibility: hidden !important;}
		[data-testid="stDecoration"] {display: none;}
	</style>
"""
st.markdown(hide_menu, unsafe_allow_html=True)


import base64





# st.set_page_config(page_title="📄→🎥 Headings Only", layout="wide")

# ─── PAGE EXTRACTION & SUMMARIZATION ─────────────────────────────────
@st.cache_data
def extract_pages_from_bytes(pdf_bytes):
	clean = lambda t: re.sub(r'\s+', ' ',
				 re.sub(r'Page \d+|Reprint 2025-26|===== Page \d+ =====', '', t)
			   ).strip()
	reader = PdfReader(pdf_bytes)
	return [clean(page.extract_text() or "") for page in reader.pages]


def summarize(pages, stopwords, target_words=TARGET_WORDS):
	text = " ".join(pages)
	words = re.findall(r'\w+', text.lower())
	freqs = Counter(w for w in words if w not in stopwords)
	sents = re.split(r'(?<=[\.\?\!])\s+', text)
	scored = [(i, sum(freqs[w] for w in re.findall(r'\w+', s.lower()))/max(len(s),1), s)
			  for i, s in enumerate(sents)]
	scored.sort(key=lambda x: x[1], reverse=True)
	sel, wc = [], 0
	for idx, _, s in scored:
		wcount = len(re.findall(r'\w+', s))
		if wc + wcount > target_words and wc > 0:
			break
		sel.append((idx, s))
		wc += wcount
	sel.sort()
	return " ".join(s for _, s in sel)

def split_into_slides(text, max_chars=380):
	chunks, cur = [], ""
	for p in text.split('. '):
		p = p.strip()
		if not p:
			continue
		p += '.' if not p.endswith('.') else ''
		if len(cur) + len(p) < max_chars:
			cur += p + " "
		else:
			chunks.append(cur.strip())
			cur = p + " "
	if cur:
		chunks.append(cur.strip())
	return chunks

# ─── HEADING EXTRACTION ─────────────────────────────────────────────────
def extract_heading(txt: str) -> str:
	"""Return only the first sentence as the slide heading."""
	# split on first period, question, or exclamation
	parts = re.split(r'(?<=[\.\?\!])\s+', txt, maxsplit=1)
	return parts[0].strip()

# ─── IMAGE SEARCH & DOWNLOAD (with caching) ────────────────────────────
@functools.lru_cache(maxsize=256)
def fetch_image_for_slide(text, prefix, idx):
	# img_path = os.path.join(IMG_DIR, f"{tag}_slide_{idx}.jpg")
	img_path = os.path.join(IMG_DIR, f"{prefix}_slide_{idx}.jpg")
	if os.path.exists(img_path):
		return img_path
	if not UNSPLASH_ACCESS_KEY:
		return None

	query = text if len(text) < 50 else text[:50]
	url = "https://api.unsplash.com/search/photos"
	params = {"query": query, "per_page": 1, "orientation": "landscape"}
	headers = {"Authorization": f"Client-ID {UNSPLASH_ACCESS_KEY}"}
	try:
		r = requests.get(url, params=params, headers=headers, timeout=5)
		r.raise_for_status()
		results = r.json().get("results")
		if not results:
			return None
		img_url = results[0]["urls"]["regular"]
		resp = requests.get(img_url, stream=True, timeout=5)
		resp.raise_for_status()
		with open(img_path, "wb") as out:
			for chunk in resp.iter_content(1024):
				out.write(chunk)
		return img_path
	except Exception:
		return img_path if os.path.exists(img_path) else None

# ─── at the top of your script, after imports ────────────────────────
VOICE_OPTIONS = {
	
	 "Indian English (co.in)": ("en", "co.in"),
	"US English (com)":        ("en", "com"  ),
	"British English (co.uk)": ("en", "co.uk"),
	# …etc…

	# Indian languages — no TLD override needed
	"Hindi (hi)":    ("hi", None),
	"Telugu (te)":   ("te", None),
	"Kannada(kn)":   ("kn",None),
	"Tamil(ta)":     ("ta",None),
	"Malayalam(ml)": ("ml",None),
	"Marathi(mr)":   ("mr",None),
	"Gujarati(gu)":  ("gu",None),
	"Bengali(bn)":   ("bn",None)
}




# ─── Load Bootstrap CSS ─────────────────────────────────────────────────
# Load Bootstrap CSS once
st.markdown("""
<link
  href="https://cdn.jsdelivr.net/npm/bootstrap@5.0.2/dist/css/bootstrap.min.css"
  rel="stylesheet"
  integrity="sha384-EVSTQN3/azprG1Anm3QDgpJLIm9Nao0Yz1ztcQTwFspd3yD65VohhpuuCOmLASjC"
  crossorigin="anonymous"
/>
<style>
#root{
	background: #2b296a;
	height
}
</style>
""", unsafe_allow_html=True)

# Make a 2-column layout: col1 for the card, col2 for the selectbox
col1, col2, col3 = st.columns([3, 6, 3])


 
with col2:
	st.markdown("""
	<style>
	.upload-container {
		text-align: center;
		# background-color: #f6f6f6;
		padding: 10px	; 
		box-shadow: 0 4px 12px rgba(0,0,0,0.05);
		font-family: 'Segoe UI', sans-serif;
		 
	}

	.upload-title {
		font-size: 32px;
		font-weight: 700;
		margin-bottom: 8px;
		color: #1a1a1a;
	} 
	 
	/* Force center align and style for the file uploader */
	section[data-testid="stFileUploader"] {
		display: flex;
		justify-content: center;
		margin-top: -20px;
		margin-bottom: 10px;
	}

	section[data-testid="stFileUploader"] label {
		background-color: #d32f2f;
		color: white !important;
		padding: 14px 28px;
		font-size: 18px;
		font-weight: bold;
		border-radius: 50px;
		cursor: pointer;
		transition: background-color 0.3s ease;
	}

	section[data-testid="stFileUploader"] label:hover {
		background-color: #b71c1c;
	}
	</style>

	<div class="upload-container">
		<div class="upload-title">Multilingual Video Generator</div>
			
	</div>
	""", unsafe_allow_html=True)

	# Appears just below the container and matches visually

 

 

# ─── TTS & SLIDE CREATION ─────────────────────────────────────────────
def synthesize_slide_gtts(txt, filename, lang, tld=None):
	if tld:
		tts = gTTS(text=txt, lang=lang, tld=tld)
	else:
		tts = gTTS(text=txt, lang=lang)
	tts.save(filename)
	return filename





def synthesize_slides(slides, prefix, lang, tld):
	files = []
	to_say = []

	# include lang/tld in your filename so different languages never collide
	tag = f"{lang}" + (f"_{tld}" if tld else "")
	for i, txt in enumerate(slides):
		out_path = os.path.join(
			AUDIO_DIR,
			f"{prefix}_{tag}_scene_{i:03d}.mp3"
		)
		files.append(out_path)
		if not os.path.exists(out_path) or os.path.getsize(out_path) == 0:
			to_say.append((txt, out_path))

	for txt, out_path in to_say:
		synthesize_slide_gtts(txt, out_path, lang, tld)

	return files





def make_slide(txt, duration, bg_image_path=None):
	W, H = 640, 360

	# pick background
	if bg_image_path and os.path.exists(bg_image_path):
		bg = Image.open(bg_image_path).convert("RGB").resize((W, H), Image.LANCZOS)
	else:
		bg = Image.new("RGB", (W, H), (30, 30, 70))

	draw = ImageDraw.Draw(bg)
	font_path = "arial.ttf"
	if selected_lang == "hi" and os.path.exists("NotoSansDevanagari-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansDevanagari-Regular.ttf", 20)

	elif selected_lang == "te":
		font_path = r"C:\Windows\Fonts\Nirmala.ttf"
		font = ImageFont.truetype(font_path, 20) if os.path.exists(font_path) else ImageFont.load_default()
	elif selected_lang == "kn" and os.path.exists("NotoSansKannada-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansKannada-Regular.ttf", 20)
	elif selected_lang == "ta" and os.path.exists("NotoSansTamil-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansTamil-Regular.ttf", 20)
	elif selected_lang == "ml" and os.path.exists("NotoSansMalayalam-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansMalayalam-Regular.ttf", 20)
	elif selected_lang == "mr" and os.path.exists("NotoSansDevanagari-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansDevanagari-Regular.ttf", 20)
	elif selected_lang == "gu" and os.path.exists("NotoSansGujarati-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansGujarati-Regular.ttf", 20)
	elif selected_lang == "bn" and os.path.exists("NotoSansBengali-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansBengali-Regular.ttf", 20)
	
	else:
	   try:
		   font = ImageFont.truetype("arial.ttf", 20)
	   except:
		   font = ImageFont.load_default()
	# only show heading
	heading = extract_heading(txt)
	lines = textwrap.wrap(heading, width=40)

	# draw at fixed top margin
	y = 20
	for line in lines:
		bbox = draw.textbbox((0, 0), line, font=font)
		w = bbox[2] - bbox[0]
		draw.text(((W - w)//2, y), line, font=font, fill="white")
		y += bbox[3] - bbox[1] + 8   # line height + small spacing

	return ImageClip(np.array(bg)).set_duration(duration)



# ─── SRT WRITING ───────────────────────────────────────────────────────
def format_timestamp(seconds):
	ms = int((seconds - int(seconds)) * 1000)
	h = int(seconds // 3600)
	m = int((seconds % 3600) // 60)
	s = int(seconds % 60)
	return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

def write_srt(slides, audio_files, prefix):
	subs = pysrt.SubRipFile()
	current = 0.0
	index = 1

	for slide_text, wav in zip(slides, audio_files):
		# load duration
		clip = AudioFileClip(wav)
		dur  = clip.duration
		clip.close()

		# split into sentences
		sents = re.split(r'(?<=[\.!?])\s+', slide_text.strip())
		if not sents:
			continue

		per = dur / len(sents)
		for i, sentence in enumerate(sents):
			start = current + i * per
			end   = start + per
			subs.append(pysrt.SubRipItem(
				index=index,
				start=pysrt.SubRipTime(milliseconds=int(start*1000)),
				end=  pysrt.SubRipTime(milliseconds=int(end*1000)),
				text= sentence.strip()
			))
			index += 1

		current += dur

	out = os.path.join(SRT_DIR, f"{prefix}.srt")
	subs.save(out, encoding="utf-8")
	return out



# ─── 2) BURN with PIL so each sentence appears/fades in sync ────────────
def burn_subtitles_pil(video_path, srt_path, out_path):
	video = VideoFileClip(video_path)
	W, H   = video.size
	fps    = video.fps

	if selected_lang == "hi" and os.path.exists("NotoSansDevanagari-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansDevanagari-Regular.ttf", 20)
	

	elif selected_lang == "te":
		font_path = r"C:\Windows\Fonts\Nirmala.ttf"
		font = ImageFont.truetype(font_path, 20) if os.path.exists(font_path) else ImageFont.load_default()
	elif selected_lang == "kn" and os.path.exists("NotoSansKannada-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansKannada-Regular.ttf", 20)
	elif selected_lang == "ta" and os.path.exists("NotoSansTamil-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansTamil-Regular.ttf", 20)
	elif selected_lang == "ml" and os.path.exists("NotoSansMalayalam-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansMalayalam-Regular.ttf", 20)
	elif selected_lang == "mr" and os.path.exists("NotoSansDevanagari-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansDevanagari-Regular.ttf", 20)
	elif selected_lang == "gu" and os.path.exists("NotoSansGujarati-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansGujarati-Regular.ttf", 20)
	elif selected_lang == "bn" and os.path.exists("NotoSansBengali-Regular.ttf"):
	   font = ImageFont.truetype("NotoSansBengali-Regular.ttf", 20)
	

	
	
	else:
	   # fallback to a system font
	   font_path = r"C:\Windows\Fonts\arial.ttf"
	   font = ImageFont.truetype(font_path, 20) if os.path.exists(font_path) else ImageFont.load_default()


	subs      = pysrt.open(srt_path)
	subtitle_clips = []
	pad_x, pad_y    = 8, 4
	margin_bot      = 30
	fade            = 0.1

	for sub in subs:
		start = sub.start.ordinal/1000.0
		end   = sub.end.ordinal/1000.0
		txt   = sub.text.replace("\n"," ")

		# wrap to max two lines
		lines = textwrap.wrap(txt, width=50)
		if len(lines) > 2:
			lines = [ lines[0], " ".join(lines[1:]) ]

		# measure text
		dummy = Image.new("RGBA", (1,1))
		d     = ImageDraw.Draw(dummy)
		widths  = [d.textbbox((0,0), ln, font=font)[2] for ln in lines]
		heights = [d.textbbox((0,0), ln, font=font)[3] for ln in lines]
		box_w = min(max(widths) + 2*pad_x, int(W*0.8))
		box_h = sum(heights) + (len(lines)+1)*pad_y

		# render subtitle image
		img = Image.new("RGBA", (box_w, box_h), (0,0,0,0))
		dd  = ImageDraw.Draw(img)
		dd.rectangle([(0,0),(box_w,box_h)], fill=(0,0,0,180))
		y = pad_y
		for ln in lines:
			w,h = d.textbbox((0,0), ln, font=font)[2:4]
			x = (box_w - w)//2
			dd.text((x, y), ln, font=font, fill="white")
			y += h + pad_y

		arr = np.array(img)
		clip = (ImageClip(arr, ismask=False)
				.set_start(start).set_end(end)
				.set_position(("center", H - box_h - margin_bot))
				.fadein(fade).fadeout(fade)
			   )
		subtitle_clips.append(clip)

	final = CompositeVideoClip([video, *subtitle_clips])
	final.write_videofile(
		out_path,
		codec="libx264", audio_codec="aac",
		fps=fps, threads=os.cpu_count()
	)





# ─── 3) GENERATE VIDEO: slides + moving subtitles ───────────────────────
def generate_video(slides, prefix,rebuild=False):

	# 1) synthesize audio
	audio_files = synthesize_slides(slides, prefix,selected_lang,selected_tld)
	st.session_state[f"audio_files_{prefix}"] = audio_files

	common_heading = extract_heading(slides[0])

	# 2) fetch images (unchanged)
	# Custom logic to support both images and uploaded video clips
	img_paths = []
	for i in range(len(slides)):
		found = False
		for ext in [".mp4", ".mov", ".jpg", ".jpeg", ".png"]:
			# path = os.path.join(IMG_DIR, f"{tag}_slide_{i}{ext}")
			path = os.path.join(IMG_DIR, f"{prefix}_slide_{i}{ext}")
			if os.path.exists(path):
				img_paths.append(path)
				found = True
				break
		if not found:
			# fallback to Unsplash image
			img_paths.append(fetch_image_for_slide(slides[i], prefix, i))


	# 3) build slide clips
	def _build(params):
		txt, aud, img_or_vid_path = params
		audio = AudioFileClip(aud)
		duration = audio.duration
		# heading = extract_heading(txt)
		heading_short = shorten(common_heading, width=60, placeholder="…")

		ext = os.path.splitext(img_or_vid_path)[-1].lower()
		if ext in ['.mp4', '.mov']:
			# Use video background
			bg_clip = VideoFileClip(img_or_vid_path).subclip(0, min(duration, VideoFileClip(img_or_vid_path).duration))
			bg_clip = bg_clip.resize((640, 360)).set_duration(duration)
		else:
			# Use image background as before
			# heading_short = shorten(heading, width=60, placeholder="…")
			bg_clip = make_slide(heading_short, duration, img_or_vid_path)

		clip = bg_clip.set_audio(AudioFileClip(aud))
		audio.close()
		return clip


	with concurrent.futures.ThreadPoolExecutor(max_workers=os.cpu_count()) as ex:
		clips = list(ex.map(_build, zip(slides, audio_files, img_paths)))

	# 4) concatenate + write raw
	raw = concatenate_videoclips(clips, method="compose")
	pdf_folder = os.path.join(VIDEO_DIR, prefix)
	os.makedirs(pdf_folder, exist_ok=True)

	tag = f"{selected_lang}" + (f"_{selected_tld}" if selected_tld else "")
	raw_filename    = f"{prefix}_{tag}.mp4"
	subtitled_fname = f"{prefix}_{tag}_subtitled.mp4"

	raw_path = os.path.join(pdf_folder, raw_filename)
	subtitled_out = os.path.join(pdf_folder, subtitled_fname)
    
    
	# write raw
	raw.write_videofile(raw_path, fps=24, codec="libx264", audio_codec="aac", 
						threads=os.cpu_count(), ffmpeg_params=["-preset","ultrafast","-crf","30"], logger=None)

	# write SRT
	srt_path = write_srt(slides, audio_files, prefix)

	# burn moving subtitles
	burn_subtitles_pil(raw_path, srt_path, subtitled_out)
	try:
		os.remove(raw_path)
	except OSError:
		pass

	return subtitled_out, srt_path


def get_slide_start_times(audio_files):
	starts, current = [], 0.0
	for wav in audio_files:
		starts.append(current)
		dur = AudioFileClip(wav).duration
		current += dur
	return starts


# uploaded = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True)

# Make a 2-column layout just like before


STOPWORDS = {"a","an","the","and","or","but","in","on","at","for","to","with"}
stopwords = STOPWORDS 

col1, col2, col3 = st.columns([3,6,3])

with col1:
	# 1) Language pick
	st.markdown("<h2 style='font-size:24px; font-weight:700'>🎙️ Language</h2>", unsafe_allow_html=True)
	voice_tld = st.selectbox("Choose Language", options=list(VOICE_OPTIONS.keys()))
	selected_lang, selected_tld = VOICE_OPTIONS[voice_tld]


	
	if "last_lang" not in st.session_state:
		st.session_state["last_lang"] = None
	if st.session_state["last_lang"] != selected_lang:
		# For each prefix with slides, regenerate translations
		for key in list(st.session_state.keys()):
			if key.startswith("slides_"):
				prefix = key.split("_", 1)[1]
				eng_slides = st.session_state[f"slides_{prefix}"]
				# Translate each slide
				translated = [
					translator.translate(text, dest=selected_lang).text
					for text in eng_slides
				]
				st.session_state[f"translated_slides_{prefix}"] = translated
		st.session_state["last_lang"] = selected_lang

	# 2) Upload PDF
	st.markdown("<h2  style='font-size:24px; font-weight:700'>📄 Upload PDF</h2>", unsafe_allow_html=True)
	uploaded_files = st.file_uploader("Upload file", type="pdf", accept_multiple_files=True, label_visibility="collapsed")


	# 2a) As soon as the PDF arrives, generate slides once:
	if uploaded_files:
		pdf = uploaded_files[0]
		prefix = os.path.splitext(pdf.name)[0]
		if f"slides_{prefix}" not in st.session_state:
			pages  = extract_pages_from_bytes(pdf)
			summary= summarize(pages, STOPWORDS)
			slides = split_into_slides(summary)
			st.session_state[f"slides_{prefix}"]    = slides
			st.session_state[f"generated_{prefix}"]  = True  # mark that slides exist

	

def rebuild_video_with_edits(slides, prefix):
	st.info("Re-building video with your edits…")
	fetch_image_for_slide.cache_clear()
	out_vid, out_srt = generate_video(slides, prefix, rebuild=True)
	st.session_state[f"video_path_{prefix}_{tag}"] = out_vid
	st.session_state[f"generated_{prefix}_{tag}"]  = True
	st.success("Re-build complete!")


# # hard-coded stopword set; no widget shown
STOPWORDS = {"a","an","the","and","or","but","in","on","at","for","to","with"}
stopwords = STOPWORDS 


if uploaded_files:

	for pdf_file in uploaded_files:
		key    = pdf_file.name
		prefix = key.replace(" ", "_").rsplit(".", 1)[0]

		# 1) generate slides on first upload
		if f"slides_{prefix}" not in st.session_state:
			pages   = extract_pages_from_bytes(pdf_file)
			summary = summarize(pages, STOPWORDS)
			slides  = split_into_slides(summary)
			st.session_state[f"slides_{prefix}"]   = slides
			# mark as "not yet generated"
			st.session_state[f"generated_{prefix}"] = False
			

		
		else:
			slides_to_display = st.session_state[f"slides_{prefix}"]

		# ─── REPLACE MEDIA UI ────────────────────────────────────────────────
		with col1:
			st.markdown("<h2  style='font-size:24px; font-weight:700'	>✏️ Replace Image/Video or Audio</h2>", unsafe_allow_html=True)
			slides = st.session_state[f"slides_{prefix}"]

			chosen     = st.selectbox("Slide # to replace", list(range(len(slides))))
			media_type = st.radio("Replace", ["Image/Video", "Audio"], horizontal=True)



			if media_type == "Image/Video":
				up = st.file_uploader(
					f"Upload new image/video for slide {chosen}",
					type=["png", "jpg", "jpeg", "mp4", "mov"],
					key=f"mediaedit_{prefix}_{chosen}"
				)
				if up:
					ext  = up.name.split(".")[-1]
					dest = os.path.join(IMG_DIR, f"{prefix}_slide_{chosen}.{ext}")
					# delete old, save new, clear cache…
					for old_ext in [".png", ".jpg", ".jpeg", ".mp4", ".mov"]:
						old_file = os.path.join(IMG_DIR, f"{prefix}_slide_{chosen}{old_ext}")
						if os.path.exists(old_file):
							os.remove(old_file)
					with open(dest, "wb") as f:
						f.write(up.read())
					fetch_image_for_slide.cache_clear()
					st.success("✅ Image/Video replaced!")

			else:  # Audio
				up = st.file_uploader(
					f"Upload new audio for slide {chosen}",
					type=["mp3"],
					key=f"audioedit_{prefix}_{chosen}"
				)
				if up:
					tag = f"{selected_lang}" + (f"_{selected_tld}" if selected_tld else "")
					dest = os.path.join(AUDIO_DIR, f"{prefix}_{tag}_scene_{chosen:03d}.mp3")
					with open(dest, "wb") as f:
						f.write(up.read())
					st.success("✅ Audio replaced!")

			# Rebuild button writes to same key as initial generate
			if st.button("Rebuild video with edits", key=f"rebuild_{prefix}"):
				out_vid, out_srt = generate_video(slides, prefix, rebuild=True)
				st.session_state[f"video_path_{prefix}"]   = out_vid
				st.session_state[f"generated_{prefix}"]    = True
				st.success("🎉 Video rebuilt!")

		# ─── EDIT TRANSCRIPT ─────────────────────────────────────────────────
		with col3:
			audio_files = st.session_state.get(f"audio_files_{prefix}", [])
			# if audio_files:
			# 	starts = get_slide_start_times(audio_files)
			# 	start_ts = format_timestamp(starts[chosen])
			# 	st.markdown(f"**Slide {chosen} starts at {start_ts}**")
			st.header("📝 Edit Transcript")
			# orig   = st.session_state[f"slides_{prefix}"]
			source_key = (
				f"translated_slides_{prefix}"
				if f"translated_slides_{prefix}" in st.session_state
				else f"slides_{prefix}"
			)
			orig = st.session_state[source_key]
			edited = st.text_area(f"Slides for {key}", "\n\n".join(orig), height=300)
			if st.button("Update Transcript", key=f"upd_{prefix}"):
				new_slides = [s.strip() for s in edited.split("\n\n") if s.strip()]
				st.session_state[f"slides_{prefix}"]  = new_slides
				st.session_state[f"generated_{prefix}"] = False
				st.success("Transcript updated!")
			

			
			
			

		# ─── VIDEO GENERATION & DISPLAY ────────────────────────────────────
		with col2:
			st.header("🎞️ Generate Video")
			# first-time generate
			if not st.session_state[f"generated_{prefix}"]:
				if st.button("Generate Video", key=f"gen_{prefix}"):
					# clean up old audio
					for fn in os.listdir(AUDIO_DIR):
						if fn.startswith(f"{prefix}_") and fn.endswith(".mp3"):
							try: os.remove(os.path.join(AUDIO_DIR, fn))
							except: pass

					with st.spinner("🔧 Building video…"):
						vid, srt = generate_video(st.session_state[f"slides_{prefix}"], prefix)
						# write to the same key as rebuild
						st.session_state[f"video_path_{prefix}"]   = vid
						st.session_state[f"generated_{prefix}"]    = True
						st.success("✅ Video ready!")

			# and in all cases, if a video_path exists, show it:
			if st.session_state.get(f"video_path_{prefix}"):
				st.video(st.session_state[f"video_path_{prefix}"])
